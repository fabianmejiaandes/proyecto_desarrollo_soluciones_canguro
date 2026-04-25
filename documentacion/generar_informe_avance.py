# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "Informe_Avance_Fundacion_Canguro.docx"
ASSETS = ROOT / "documentacion" / "informe_assets"

fig_missing = ASSETS / "EDA_Prematurez_cell04_out01.png"
fig_auc = ASSETS / "Modelado_Malnutricion-final_cell10_out00.png"
fig_shap = ASSETS / "Modelado_Malnutricion-final_cell18_out00.png"

auc_rows = [
    ("F0 Prenatal/parto", "0,645", "0,618", "0,555"),
    ("F1 Nacimiento", "0,737", "0,751", "0,689"),
    ("F2 Hospitalización", "0,741", "0,754", "0,705"),
    ("F3 40 semanas EC", "0,768", "0,773", "0,725"),
    ("F4 3 meses EC", "0,821", "0,874", "0,826"),
    ("F5 6 meses EC", "0,894", "0,936", "0,895"),
    ("F6 9 meses EC", "0,929", "0,963", "0,925"),
]

test_rows = [
    ("Stunting", "0,923", "0,818", "0,848", "6.191", "1.525"),
    ("Bajo peso", "0,959", "0,834", "0,939", "5.951", "700"),
    ("Wasting", "0,911", "0,702", "0,941", "5.950", "265"),
]

phase_rows = [
    ("F0", "Prenatal / parto", "41"),
    ("F1", "Nacimiento", "75"),
    ("F2", "Hospitalización", "107"),
    ("F3", "40 semanas EC", "137"),
    ("F4", "3 meses EC", "164"),
    ("F5", "6 meses EC", "183"),
    ("F6", "9 meses EC", "198"),
]

shap_rows = [
    ("zscoretalla9", "Talla para la edad a los 9 meses EC", "1,324"),
    ("zscoretalla6", "Talla para la edad a los 6 meses EC", "0,695"),
    ("zscorepeso9", "Peso para la edad a los 9 meses EC", "0,153"),
    ("velocidad9_6mesesOMS", "Velocidad de crecimiento 6-9 meses EC", "0,112"),
    ("zscoretalla9cat", "Categoría clínica del z-score de talla a 9 meses", "0,093"),
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color="BFBFBF"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def style_table(table, header_fill="D9EAF7"):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for r_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                for run in p.runs:
                    run.font.size = Pt(8.5)
            if r_idx == 0:
                set_cell_shading(cell, header_fill)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(31, 78, 121)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(6 if level == 1 else 4)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_para(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(1.5)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.add_run(item)


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(89, 89, 89)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    style_table(table)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    return table


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(9.5)
    styles["Normal"].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    styles["Normal"].paragraph_format.space_after = Pt(3)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[style_name].font.name = "Calibri"
        styles[style_name].font.color.rgb = RGBColor(31, 78, 121)
    styles["Heading 1"].font.size = Pt(13)
    styles["Heading 2"].font.size = Pt(11)
    styles["Heading 3"].font.size = Pt(10)
    styles["List Bullet"].font.name = "Calibri"
    styles["List Bullet"].font.size = Pt(9.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Informe de avance del proyecto de grado")
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(31, 78, 121)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "Identificación de factores de riesgo de malnutrición en niños prematuros "
        "evaluados a 12 meses de edad corregida"
    )
    r.bold = True
    r.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    p.add_run("Autores: ________________________________________________     Fecha: abril de 2026")

    add_heading(doc, "1. Descripción breve de la propuesta de solución", 1)
    add_para(
        doc,
        "El proyecto busca apoyar la toma de decisiones clínicas del Programa Madre Canguro Integral "
        "(PMCI) mediante modelos de aprendizaje automático que estimen el riesgo de malnutrición a "
        "los 12 meses de edad corregida (EC) en niños prematuros o con bajo peso al nacer. La propuesta "
        "parte de una necesidad clínica concreta: identificar tempranamente factores de riesgo modificables "
        "y no modificables para priorizar intervenciones nutricionales y de seguimiento antes de que el "
        "déficit sea evidente a los 12 meses."
    )
    add_para(
        doc,
        "La solución desarrollada hasta el momento es una cascada temporal de modelos LightGBM que calcula "
        "el riesgo en siete momentos clínicos acumulativos, desde variables prenatales y del parto hasta la "
        "visita de 9 meses EC. Los desenlaces se definieron con criterios antropométricos usados internacionalmente: "
        "stunting o talla baja (HAZ < -2 DS), bajo peso (WAZ < -2 DS) y wasting o desnutrición aguda "
        "(WHZ < -2 DS), coherentes con los estándares de crecimiento de la OMS (World Health Organization, 2006)."
    )
    add_para(
        doc,
        "El enfoque mantiene una lectura clínica: no solo produce probabilidades, sino que identifica variables "
        "influyentes mediante SHAP, lo que permite discutir los resultados con neonatólogos y convertir el modelo "
        "en una herramienta de soporte, no en un sustituto del juicio médico."
    )

    add_heading(doc, "2. Recolección y preparación de los datos", 1)
    add_para(
        doc,
        "La fuente principal es la base clínica anonimizada del PMCI de la Fundación Canguro, exportada desde "
        "SPSS a Excel. El archivo contiene 64.801 historias clínicas y 753 variables recolectadas longitudinalmente "
        "durante el seguimiento: prenatal, nacimiento, hospitalización, 40 semanas de edad gestacional, 3, 6, 9 "
        "y 12 meses EC. El diccionario PhETI documenta 634 variables y permite relacionar variables con fases, "
        "episodios y temas de interés."
    )
    add_para(
        doc,
        "Durante la preparación se transformaron los valores \"#NULL!\" a valores faltantes reales, se convirtieron "
        "a numéricas las columnas con al menos 50% de valores numéricos válidos y se construyeron tres variables "
        "objetivo binarias a partir de las categorías de z-score de los 12 meses. Además, se excluyeron variables "
        "con fuga de información, en especial columnas posteriores a 9 meses o derivadas directamente del resultado "
        "de 12 meses, como variables con \"12\" en el nombre, salvo los outcomes usados para etiquetar el desenlace."
    )

    add_table(doc, ["Fase", "Momento clínico", "Variables acumuladas"], phase_rows, widths=[0.7, 3.1, 1.4])
    add_caption(doc, "Tabla 1. Cascada temporal de variables utilizada para entrenamiento e inferencia.")

    if fig_missing.exists():
        doc.add_picture(str(fig_missing), width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_caption(doc, "Figura 1. Distribución de valores faltantes por columna en el EDA inicial.")

    add_para(
        doc,
        "El EDA evidenció una calidad heterogénea: 1 columna sin nulos, 442 columnas con 1-50% de nulos, "
        "226 columnas con 51-75% y 84 columnas con más de 75%. En la versión actual del pipeline, LightGBM "
        "aprovecha su manejo nativo de valores faltantes y evita imputaciones globales que podrían introducir "
        "sesgo. Para análisis temporales se priorizaron los periodos P4-P6, correspondientes a 2007-2012, "
        "2013-2017 y 2018-2022, debido a su mayor completitud; los registros de 2023 se consideraron solo "
        "como sensibilidad por seguimiento incompleto."
    )

    add_heading(doc, "3. Construcción de los modelos o solución", 1)
    add_para(
        doc,
        "Se entrenaron 21 modelos binarios, uno por cada combinación de fase y outcome (7 fases x 3 desenlaces). "
        "La arquitectura temporal permite responder dos preguntas complementarias: qué tan temprano puede estimarse "
        "el riesgo y cuánto mejora la predicción al incorporar mediciones sucesivas del seguimiento clínico."
    )
    add_para(
        doc,
        "El algoritmo seleccionado fue LightGBM por su desempeño en datos tabulares, su tolerancia a valores faltantes "
        "y su capacidad para capturar relaciones no lineales. Cada entrenamiento usó validación cruzada estratificada "
        "de 5 folds; el desbalance de clases se compensó con scale_pos_weight calculado por fold. Los hiperparámetros "
        "principales fueron learning_rate = 0,05, num_leaves = 63, min_child_samples = 30, feature_fraction = 0,8, "
        "bagging_fraction = 0,8 y early stopping de 50 rondas."
    )
    add_para(
        doc,
        "Después de la validación, se entrenaron modelos finales sobre el 100% de los datos disponibles con el número "
        "promedio de iteraciones observado en validación cruzada. Los modelos quedaron persistidos en formato nativo "
        "LightGBM (.lgb), acompañados de archivos JSON con la lista de variables esperadas por fase. También se generaron "
        "archivos para tablero: predicciones de test, métricas por fase, valores SHAP y documentación de lectura."
    )
    add_para(
        doc,
        "Como referencia interpretable, se comparó el mejor modelo de stunting en F6 contra regresión logística L1. "
        "LightGBM obtuvo AUC 0,9290 frente a 0,9216 +/- 0,0031 de la regresión logística, una ganancia moderada que "
        "sugiere que parte importante de la señal clínica es lineal, aunque LightGBM conserva ventajas por manejo de "
        "no linealidades, faltantes e interacción entre variables."
    )

    add_heading(doc, "4. Resultados obtenidos", 1)
    add_para(
        doc,
        "La señal predictiva aumenta de forma consistente a medida que se agregan fases del seguimiento. En validación "
        "cruzada, el AUC de stunting pasa de 0,645 en F0 a 0,929 en F6; bajo peso llega a 0,963 y wasting a 0,925. "
        "El mayor salto ocurre al incorporar información de 3, 6 y 9 meses EC, lo que confirma que la trayectoria de "
        "crecimiento es central para anticipar el estado nutricional a los 12 meses."
    )

    add_table(doc, ["Fase", "Stunting", "Bajo peso", "Wasting"], auc_rows, widths=[2.2, 1.0, 1.0, 1.0])
    add_caption(doc, "Tabla 2. ROC-AUC media por fase y outcome en validación cruzada estratificada de 5 folds.")

    if fig_auc.exists():
        doc.add_picture(str(fig_auc), width=Inches(6.4))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_caption(doc, "Figura 2. Evolución del AUC a lo largo de la cascada temporal.")

    add_para(
        doc,
        "En el conjunto de test 80/20 estratificado, la mejor fase fue F6 para los tres desenlaces. Bajo peso alcanzó "
        "el desempeño más alto, con AUC 0,959, sensibilidad 0,834 y especificidad 0,939. Stunting logró AUC 0,923, "
        "sensibilidad 0,818 y especificidad 0,848. Wasting mantuvo AUC alto (0,911), pero menor sensibilidad (0,702), "
        "consistente con su menor prevalencia y con la dificultad de detectar eventos poco frecuentes sin aumentar "
        "falsos positivos."
    )
    add_table(doc, ["Outcome en F6", "AUC test", "Sens.", "Espec.", "n test", "positivos"], test_rows, widths=[1.7, 0.9, 0.8, 0.8, 0.8, 0.8])
    add_caption(doc, "Tabla 3. Métricas de test en la mejor fase disponible (F6: 9 meses EC).")

    if fig_shap.exists():
        doc.add_picture(str(fig_shap), width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_caption(doc, "Figura 3. Importancia global SHAP para stunting en la fase F6.")

    add_table(doc, ["Variable", "Lectura clínica", "SHAP"], shap_rows, widths=[1.8, 3.5, 0.7])
    add_caption(doc, "Tabla 4. Principales factores explicativos para stunting a 12 meses EC.")

    add_heading(doc, "5. Análisis de los resultados obtenidos y próximos pasos", 1)
    add_para(
        doc,
        "Los resultados son coherentes con la hipótesis clínica: el riesgo nutricional a 12 meses no depende de una "
        "medición aislada, sino de una trayectoria acumulada desde el nacimiento y el seguimiento ambulatorio. La mejora "
        "progresiva del AUC muestra que la cascada temporal es adecuada para dos usos distintos: alerta temprana en F1-F3, "
        "con desempeño moderado, y priorización más precisa en F4-F6, cuando la señal antropométrica reciente es mucho "
        "más fuerte."
    )
    add_para(
        doc,
        "La interpretabilidad refuerza esta lectura. Para stunting, los predictores principales son z-scores de talla a "
        "9 y 6 meses, peso a 9 meses y velocidad de crecimiento entre 6 y 9 meses. Esto indica que la ventana de 6 a "
        "9 meses EC es crítica: todavía está antes del desenlace de 12 meses, pero suficientemente cerca para capturar "
        "desaceleraciones que podrían activar intervención nutricional intensiva. La talla materna aparece como factor "
        "contextual no modificable, útil para ajustar expectativas clínicas sin convertirla en criterio único de riesgo."
    )
    add_para(
        doc,
        "También hay limitaciones importantes. Primero, el desempeño más alto ocurre en F6, una fase cercana al outcome, "
        "por lo que su utilidad es más de intervención tardía/preventiva antes de 12 meses que de predicción neonatal "
        "temprana. Segundo, wasting tiene pocos casos positivos; por ello se deben ajustar umbrales con criterio clínico, "
        "priorizando sensibilidad si el costo de falsos negativos es alto. Tercero, la pérdida de seguimiento y los faltantes "
        "no son aleatorios: muchas variables de 3, 6 y 9 meses dependen de continuidad en el programa. Finalmente, aunque "
        "se auditó leakage eliminando variables de 12 meses, se debe mantener una revisión clínica de variables derivadas "
        "para evitar que indicadores demasiado próximos al outcome conviertan el modelo en una medición retrospectiva."
    )
    add_para(
        doc,
        "Las cohortes temporales sugieren, además, una reducción de la prevalencia de stunting entre 2007-2012 y 2018-2022, "
        "de 27,1% a 19,5%. Esta variación puede reflejar mejoras en protocolos, cambios poblacionales o diferencias de "
        "registro; por tanto, la validación temporal debe tratarse como componente central antes de llevar la solución a "
        "una interfaz clínica."
    )

    add_bullets(
        doc,
        [
            "Validar con neonatólogos la definición de outcome compuesto y los umbrales operativos de riesgo por desenlace.",
            "Agregar evaluación de calibración (Brier score y curva de calibración) para que las probabilidades sean clínicamente interpretables.",
            "Profundizar la validación temporal entre cohortes y documentar posibles cambios de distribución de variables.",
            "Definir umbrales por uso clínico: tamizaje temprano con alta sensibilidad y priorización de recursos con mejor especificidad.",
            "Construir prototipo de tablero que consuma dashboard_data/ y permita explicar predicciones individuales con SHAP.",
            "Evaluar uno o dos casos de uso reales con equipos de neonatólogos, incorporando retroalimentación sobre usabilidad, interpretación y riesgo ético.",
        ],
    )

    add_heading(doc, "Referencias", 1)
    refs = [
        "Bhutta, Z. A., Das, J. K., Rizvi, A., et al. (2017). Evidence-based interventions for improvement of maternal and child nutrition: what can be done and at what cost? The Lancet, 389(10064), 452-477. https://doi.org/10.1016/S0140-6736(16)32410-9",
        "Charpak, N., Ruiz, J. G., Zupan, J., et al. (2005). Kangaroo mother care: 25 years after. Acta Paediatrica, 94(5), 514-522. https://doi.org/10.1111/j.1651-2227.2005.tb01930.x",
        "Charpak, N., Tessier, R., Ruiz, J. G., et al. (2017). Twenty-year follow-up of kangaroo mother care versus traditional care. Pediatrics, 139(1), e20162063. https://doi.org/10.1542/peds.2016-2063",
        "Charpak, N., & Montealegre-Pomar, A. (2023). Follow-up of Kangaroo Mother Care programmes in the last 28 years: results from a cohort of 57,154 low-birth-weight infants in Colombia. BMJ Global Health, 8(5).",
        "Fenton, T. R., & Kim, J. H. (2013). A systematic review and meta-analysis to revise the Fenton growth chart for preterm infants. BMC Pediatrics, 13, 59. https://doi.org/10.1186/1471-2431-13-59",
        "Fundación Canguro. (2026). Propuesta de proyecto: identificación de factores de riesgo de malnutrición en niños prematuros evaluados a 12 meses de edad corregida. Documento interno del proyecto.",
        "World Health Organization. (2006). WHO Child Growth Standards: length/height-for-age, weight-for-age, weight-for-length, weight-for-height and body mass index-for-age: methods and development. WHO.",
        "Yang, Q., et al. (2023). Reporting and risk of bias of prediction models based on machine learning in preterm birth: a systematic review. Documento citado en la propuesta académica del proyecto.",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        r = p.add_run(ref)
        r.font.size = Pt(8.5)

    footer = section.footer.paragraphs[0]
    footer.text = "Informe de avance - PMCI / Fundación Canguro"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(127, 127, 127)

    doc.save(OUT)
    return doc


if __name__ == "__main__":
    document = build_document()
    print(OUT)
    print(f"Documento generado con {len(document.paragraphs)} párrafos, {len(document.tables)} tablas.")
