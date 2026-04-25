# -*- coding: utf-8 -*-
"""Genera una versión del informe con lenguaje claro.

La primera versión del informe se conserva intacta. Este script reutiliza las
tablas, figuras y estilo general, pero reescribe las secciones 2 en adelante
para que una persona externa al proyecto pueda entender el avance sin dificultad.
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from pathlib import Path
import importlib.util


ROOT = Path(__file__).resolve().parents[1]
BASE_SCRIPT = ROOT / "documentacion" / "generar_informe_avance.py"
OUT = ROOT / "Informe_Avance_Fundacion_Canguro_Lenguaje_Claro.docx"

spec = importlib.util.spec_from_file_location("base_informe", BASE_SCRIPT)
base = importlib.util.module_from_spec(spec)
spec.loader.exec_module(base)


def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(9.5)
    styles["Normal"].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    styles["Normal"].paragraph_format.space_after = Pt(3)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[style_name].font.name = "Calibri"
        styles[style_name].font.color.rgb = RGBColor(31, 78, 121)
    styles["Heading 1"].font.size = Pt(13)
    styles["Heading 2"].font.size = Pt(11)
    styles["Heading 3"].font.size = Pt(10)
    styles["List Bullet"].font.name = "Calibri"
    styles["List Bullet"].font.size = Pt(9.5)
    return doc


def add_title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Informe de avance del proyecto de grado")
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(31, 78, 121)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "Identificación de factores de riesgo de malnutrición en niños prematuros "
        "evaluados a 12 meses de edad corregida"
    )
    r.bold = True
    r.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    p.add_run("Autores: ________________________________________________     Fecha: abril de 2026")


def add_section_1_unchanged(doc):
    base.add_heading(doc, "1. Descripción breve de la propuesta de solución", 1)
    base.add_para(
        doc,
        "El proyecto busca apoyar la toma de decisiones clínicas del Programa Madre Canguro Integral "
        "(PMCI) mediante modelos de aprendizaje automático que estimen el riesgo de malnutrición a "
        "los 12 meses de edad corregida (EC) en niños prematuros o con bajo peso al nacer. La propuesta "
        "parte de una necesidad clínica concreta: identificar tempranamente factores de riesgo modificables "
        "y no modificables para priorizar intervenciones nutricionales y de seguimiento antes de que el "
        "déficit sea evidente a los 12 meses."
    )
    base.add_para(
        doc,
        "La solución desarrollada hasta el momento es una cascada temporal de modelos LightGBM que calcula "
        "el riesgo en siete momentos clínicos acumulativos, desde variables prenatales y del parto hasta la "
        "visita de 9 meses EC. Los desenlaces se definieron con criterios antropométricos usados internacionalmente: "
        "stunting o talla baja (HAZ < -2 DS), bajo peso (WAZ < -2 DS) y wasting o desnutrición aguda "
        "(WHZ < -2 DS), coherentes con los estándares de crecimiento de la OMS (World Health Organization, 2006)."
    )
    base.add_para(
        doc,
        "El enfoque mantiene una lectura clínica: no solo produce probabilidades, sino que identifica variables "
        "influyentes mediante SHAP, lo que permite discutir los resultados con neonatólogos y convertir el modelo "
        "en una herramienta de soporte, no en un sustituto del juicio médico."
    )


def build_document():
    doc = setup_document()
    add_title(doc)
    add_section_1_unchanged(doc)

    base.add_heading(doc, "2. Recolección y preparación de los datos", 1)
    base.add_para(
        doc,
        "La información usada en el proyecto proviene de las historias clínicas anonimizadas del Programa "
        "Madre Canguro Integral de la Fundación Canguro. En total, la base contiene 64.801 registros de niños "
        "atendidos en el programa y 753 columnas con datos clínicos, familiares y de seguimiento. Estos datos "
        "permiten observar la evolución de cada niño desde antes del nacimiento hasta los controles de 3, 6, "
        "9 y 12 meses de edad corregida."
    )
    base.add_para(
        doc,
        "Antes de construir los modelos fue necesario ordenar y limpiar la base. Algunos campos venían marcados "
        "como \"#NULL!\", que en la práctica significa que el dato no estaba disponible; esos valores se trataron "
        "como información faltante. También se revisó cuáles columnas correspondían realmente a cada momento del "
        "seguimiento para evitar usar, por accidente, información del futuro. Por ejemplo, si el modelo pretende "
        "predecir el estado nutricional a los 12 meses, no puede usar variables medidas en esa misma visita como "
        "si fueran datos disponibles con anticipación."
    )
    base.add_para(
        doc,
        "La organización por fases permite hacer predicciones en distintos momentos de la atención. En F0 solo "
        "se usa información prenatal y del parto; luego cada fase agrega nuevos datos. Así, el modelo puede dar "
        "una alerta temprana cuando hay poca información, y una estimación más precisa cuando ya existen más "
        "controles del niño."
    )

    base.add_table(doc, ["Fase", "Momento clínico", "Variables acumuladas"], base.phase_rows, widths=[0.7, 3.1, 1.4])
    base.add_caption(doc, "Tabla 1. Momentos del seguimiento usados por la solución.")

    if base.fig_missing.exists():
        doc.add_picture(str(base.fig_missing), width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        base.add_caption(
            doc,
            "Figura 1. Resumen de cuánta información faltante existe en las columnas de la base de datos.",
        )

    base.add_para(
        doc,
        "El análisis inicial mostró que no todas las columnas tienen la misma calidad. Algunas están casi completas "
        "y otras tienen muchos datos faltantes, especialmente las relacionadas con visitas de seguimiento. Esto no "
        "siempre significa un error: en un programa clínico real, algunos pacientes dejan de asistir, llegan tarde "
        "a controles o no tienen todas las mediciones. Por esa razón, se eligió un modelo capaz de trabajar con "
        "datos incompletos sin obligar a inventar valores para todos los campos."
    )
    base.add_para(
        doc,
        "Para los análisis por tiempo se trabajó principalmente con los periodos 2007-2012, 2013-2017 y 2018-2022, "
        "porque son los grupos con información más completa y comparable. Los registros de 2023 se dejaron como "
        "revisión adicional, ya que muchos niños aún no tenían cerrado el seguimiento de 12 meses."
    )

    base.add_heading(doc, "3. Construcción de los modelos o solución", 1)
    base.add_para(
        doc,
        "La solución construida funciona como una serie de modelos que se activan según el momento del seguimiento. "
        "En lugar de crear un único modelo final, se entrenaron modelos separados para cada fase y para cada tipo "
        "de malnutrición evaluado: talla baja para la edad, bajo peso para la edad y desnutrición aguda. En total "
        "se guardaron 21 modelos principales."
    )
    base.add_para(
        doc,
        "El método usado se llama LightGBM. En términos simples, es una técnica que aprende patrones a partir de "
        "tablas grandes de datos. Resulta útil en este proyecto porque puede manejar muchas variables, detectar "
        "relaciones que no son completamente lineales y trabajar con campos faltantes. Además, permite revisar "
        "qué variables influyen más en cada predicción."
    )
    base.add_para(
        doc,
        "Para comprobar que los resultados no dependieran de una sola división de los datos, se usó una estrategia "
        "de validación en cinco partes. Esto consiste en entrenar y evaluar el modelo varias veces con subconjuntos "
        "distintos de la base. También se tuvo en cuenta que algunos desenlaces son menos frecuentes que otros; por "
        "ejemplo, la desnutrición aguda tiene menos casos positivos que la talla baja. Por eso, durante el entrenamiento "
        "se ajustó el peso de las clases para que el modelo no ignorara los casos menos comunes."
    )
    base.add_para(
        doc,
        "Una vez evaluados los modelos, se guardaron versiones finales listas para ser usadas en una futura herramienta "
        "o tablero. Junto a los modelos se generaron archivos con las predicciones, las métricas de desempeño y las "
        "variables necesarias para explicar por qué un niño queda clasificado con mayor o menor riesgo."
    )
    base.add_para(
        doc,
        "También se comparó el resultado con una alternativa más tradicional, la regresión logística. La diferencia a "
        "favor de LightGBM fue moderada, no enorme. Esto es una buena señal: significa que parte de los patrones son "
        "clínicamente consistentes y no dependen de una técnica difícil de interpretar. Aun así, LightGBM aporta ventajas "
        "prácticas para manejar datos incompletos y combinar muchas variables."
    )

    base.add_heading(doc, "4. Resultados obtenidos", 1)
    base.add_para(
        doc,
        "Los resultados muestran que el modelo mejora a medida que avanza el seguimiento del niño. Esto era esperable: "
        "al comienzo solo se conocen datos del embarazo, parto y nacimiento; más adelante ya se observa cómo crece el "
        "niño en sus controles. Por eso, las predicciones hechas con información de 6 y 9 meses son más precisas que "
        "las predicciones hechas solo con datos iniciales."
    )
    base.add_para(
        doc,
        "La tabla siguiente usa AUC, una medida que resume qué tan bien separa el modelo a los niños con riesgo de los "
        "niños sin riesgo. Un AUC cercano a 0,5 equivale a una predicción casi al azar; un valor más cercano a 1 indica "
        "mejor capacidad de diferenciación. En los tres desenlaces, los valores suben de forma clara entre F0 y F6."
    )

    base.add_table(doc, ["Fase", "Stunting", "Bajo peso", "Wasting"], base.auc_rows, widths=[2.2, 1.0, 1.0, 1.0])
    base.add_caption(doc, "Tabla 2. Desempeño promedio del modelo en cada fase. Valores más altos indican mejor separación entre riesgo y no riesgo.")

    if base.fig_auc.exists():
        doc.add_picture(str(base.fig_auc), width=Inches(6.4))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        base.add_caption(
            doc,
            "Figura 2. Mejora progresiva del desempeño cuando se agrega más información del seguimiento.",
        )

    base.add_para(
        doc,
        "En la mejor fase disponible, que corresponde a los 9 meses de edad corregida, el modelo tuvo un desempeño "
        "alto para los tres desenlaces. Para bajo peso alcanzó el mejor resultado global. Para talla baja también "
        "logró una buena combinación entre detectar casos de riesgo y evitar falsas alarmas. En desnutrición aguda "
        "el resultado fue bueno, pero la sensibilidad fue menor, probablemente porque hay menos casos positivos y "
        "eso hace más difícil aprender el patrón."
    )
    base.add_para(
        doc,
        "En esta tabla, sensibilidad significa la capacidad de detectar a los niños que sí presentan el problema, "
        "mientras que especificidad significa la capacidad de reconocer a quienes no lo presentan. En un contexto "
        "clínico, estas dos medidas deben balancearse con cuidado: si se busca tamizaje temprano, puede ser preferible "
        "aceptar más alertas a cambio de dejar menos casos sin detectar."
    )
    base.add_table(doc, ["Outcome en F6", "AUC test", "Sens.", "Espec.", "n test", "positivos"], base.test_rows, widths=[1.7, 0.9, 0.8, 0.8, 0.8, 0.8])
    base.add_caption(doc, "Tabla 3. Resultados en el conjunto de prueba usando la información disponible a los 9 meses EC.")

    if base.fig_shap.exists():
        doc.add_picture(str(base.fig_shap), width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        base.add_caption(
            doc,
            "Figura 3. Variables que más ayudan a explicar la predicción de talla baja a los 12 meses.",
        )

    base.add_para(
        doc,
        "La revisión de variables importantes ayuda a entender qué está aprendiendo el modelo. Para talla baja, los "
        "factores más fuertes son mediciones recientes de crecimiento: talla a los 9 meses, talla a los 6 meses, peso "
        "a los 9 meses y velocidad de crecimiento entre 6 y 9 meses. Esto sugiere que el seguimiento del crecimiento "
        "lineal es clave para anticipar el estado nutricional a los 12 meses."
    )
    base.add_table(doc, ["Variable", "Lectura clínica", "SHAP"], base.shap_rows, widths=[1.8, 3.5, 0.7])
    base.add_caption(doc, "Tabla 4. Variables con mayor peso explicativo para la predicción de talla baja.")

    base.add_heading(doc, "5. Análisis de los resultados obtenidos y próximos pasos", 1)
    base.add_para(
        doc,
        "En conjunto, los resultados apoyan la idea central del proyecto: el riesgo de malnutrición a los 12 meses "
        "se va construyendo a lo largo del seguimiento y no depende de un solo dato aislado. El modelo puede aportar "
        "alertas tempranas desde las primeras fases, aunque con menor precisión, y puede entregar estimaciones más "
        "confiables cuando ya existen mediciones de crecimiento de 3, 6 y 9 meses."
    )
    base.add_para(
        doc,
        "La principal lectura clínica es que la ventana entre 6 y 9 meses de edad corregida parece especialmente "
        "importante. Si un niño muestra baja talla para la edad o desaceleración del crecimiento en ese periodo, el "
        "modelo lo identifica como una señal relevante de riesgo posterior. Esto puede ayudar a priorizar intervenciones "
        "nutricionales antes de llegar al control de 12 meses."
    )
    base.add_para(
        doc,
        "Sin embargo, el avance también muestra límites que deben tenerse en cuenta. El mejor desempeño aparece cerca "
        "del desenlace final, por lo que aún es necesario fortalecer la utilidad de las fases más tempranas. Además, "
        "los datos faltantes pueden estar relacionados con abandono del seguimiento, dificultades de acceso o cambios "
        "en la forma de registrar información. Estos factores pueden influir en el modelo y deben discutirse con el "
        "equipo clínico."
    )
    base.add_para(
        doc,
        "Otro punto importante es que no todos los errores tienen el mismo costo. En salud, dejar de detectar a un niño "
        "en riesgo puede ser más grave que generar una alerta adicional. Por eso, los umbrales de clasificación no deben "
        "definirse solo con criterios matemáticos, sino con apoyo de neonatólogos y nutricionistas del PMCI."
    )
    base.add_para(
        doc,
        "El análisis por cohortes también es relevante. La prevalencia de talla baja bajó de 27,1% en 2007-2012 a 19,5% "
        "en 2018-2022. Esta diferencia puede reflejar mejoras del programa, cambios en la población atendida o cambios "
        "en el registro. Antes de usar el modelo en una herramienta clínica, conviene validar que funcione bien en "
        "distintos periodos y sedes."
    )

    base.add_bullets(
        doc,
        [
            "Revisar con neonatólogos si la definición de los grupos nutricionales refleja adecuadamente la práctica clínica.",
            "Definir umbrales de alerta según el uso esperado: tamizaje temprano, priorización de seguimiento o apoyo a intervención nutricional.",
            "Comprobar que las probabilidades del modelo sean confiables y no solo sirvan para ordenar pacientes por riesgo.",
            "Evaluar con más detalle si el modelo se comporta igual en diferentes periodos de atención.",
            "Construir un tablero sencillo que muestre el riesgo del paciente y las razones principales de la alerta.",
            "Probar uno o dos casos de uso con equipos clínicos para recoger retroalimentación sobre utilidad, claridad y posibles riesgos.",
        ],
    )

    base.add_heading(doc, "Referencias", 1)
    refs = [
        "Bhutta, Z. A., Das, J. K., Rizvi, A., et al. (2017). Evidence-based interventions for improvement of maternal and child nutrition: what can be done and at what cost? The Lancet, 389(10064), 452-477. https://doi.org/10.1016/S0140-6736(16)32410-9",
        "Charpak, N., Ruiz, J. G., Zupan, J., et al. (2005). Kangaroo mother care: 25 years after. Acta Paediatrica, 94(5), 514-522. https://doi.org/10.1111/j.1651-2227.2005.tb01930.x",
        "Charpak, N., Tessier, R., Ruiz, J. G., et al. (2017). Twenty-year follow-up of kangaroo mother care versus traditional care. Pediatrics, 139(1), e20162063. https://doi.org/10.1542/peds.2016-2063",
        "Charpak, N., & Montealegre-Pomar, A. (2023). Follow-up of Kangaroo Mother Care programmes in the last 28 years: results from a cohort of 57,154 low-birth-weight infants in Colombia. BMJ Global Health, 8(5).",
        "Fenton, T. R., & Kim, J. H. (2013). A systematic review and meta-analysis to revise the Fenton growth chart for preterm infants. BMC Pediatrics, 13, 59. https://doi.org/10.1186/1471-2431-13-59",
        "Fundación Canguro. (2026). Propuesta de proyecto: identificación de factores de riesgo de malnutrición en niños prematuros evaluados a 12 meses de edad corregida. Documento interno del proyecto.",
        "World Health Organization. (2006). WHO Child Growth Standards: length/height-for-age, weight-for-age, weight-for-length, weight-for-height and body mass index-for-age: methods and development. WHO.",
        "Yang, Q., et al. (2023). Reporting and risk of bias of prediction models based on machine learning in preterm birth: a systematic review. Documento citado en la propuesta académica del proyecto.",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        r = p.add_run(ref)
        r.font.size = Pt(8.5)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.text = "Informe de avance - PMCI / Fundación Canguro - Versión en lenguaje claro"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(127, 127, 127)

    doc.save(OUT)
    return doc


if __name__ == "__main__":
    document = build_document()
    print(OUT)
    print(f"Documento generado con {len(document.paragraphs)} párrafos, {len(document.tables)} tablas.")
