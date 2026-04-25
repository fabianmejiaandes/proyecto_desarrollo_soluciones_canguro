# -*- coding: utf-8 -*-
"""Genera documentación técnica completa del proyecto PMCI/Fundación Canguro.

Este documento no reemplaza los informes cortos ya generados. Su propósito es
servir como manual de entendimiento para un programador con nociones básicas de
inteligencia artificial que necesite comprender el proyecto de extremo a extremo.
"""
from __future__ import annotations

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from pathlib import Path
import importlib.util


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "Documentacion_Tecnica_Completa_Fundacion_Canguro.docx"
BASE_SCRIPT = ROOT / "documentacion" / "generar_informe_avance.py"
ASSETS = ROOT / "documentacion" / "informe_assets"

spec = importlib.util.spec_from_file_location("base_informe", BASE_SCRIPT)
base = importlib.util.module_from_spec(spec)
spec.loader.exec_module(base)


IMAGES = {
    "preterm_distribution": ASSETS / "EDA_Prematurez_cell02_out01.png",
    "missing_values": ASSETS / "EDA_Prematurez_cell04_out01.png",
    "prematurity_key_vars": ASSETS / "EDA_Prematurez_cell06_out00.png",
    "missing_strategy": ASSETS / "EDA_Prematurez_cell08_out01.png",
    "numeric_distributions": ASSETS / "EDA_Prematurez_cell10_out00.png",
    "auc_cascade": ASSETS / "Modelado_Malnutricion-final_cell10_out00.png",
    "stunting_metrics": ASSETS / "Modelado_Malnutricion-final_cell11_out00.png",
    "roc_curves": ASSETS / "Modelado_Malnutricion-final_cell12_out00.png",
    "pr_curves": ASSETS / "Modelado_Malnutricion-final_cell13_out00.png",
    "shap_summary": ASSETS / "Modelado_Malnutricion-final_cell17_out00.png",
    "shap_bar": ASSETS / "Modelado_Malnutricion-final_cell18_out00.png",
    "dynamic_risk": ASSETS / "Modelado_Malnutricion-final_cell21_out00.png",
    "cohort_analysis": ASSETS / "Modelado_Malnutricion-final_cell30_out00.png",
    "cross_cohort": ASSETS / "Modelado_Malnutricion-final_cell31_out01.png",
    "feature_drift": ASSETS / "Modelado_Malnutricion-final_cell32_out01.png",
    "inference_demo": ASSETS / "Modelado_Malnutricion-final_cell38_out01.png",
}


phase_rows = [
    ("F0_Prenatal_Parto", "Prenatal y parto", "41", "Variables maternas, sociales, embarazo y parto."),
    ("F1_Nacimiento", "Nacimiento", "75", "Agrega peso, talla, sexo, edad gestacional, RCIU y variables del recién nacido."),
    ("F2_Hospitalizacion", "Hospitalización", "107", "Agrega días de hospitalización, oxígeno, UCI, complicaciones y alimentación hospitalaria."),
    ("F3_40semanas", "40 semanas EC", "137", "Agrega mediciones de la primera visita ambulatoria o equivalente a término."),
    ("F4_3meses", "3 meses EC", "164", "Agrega controles de crecimiento y alimentación a 3 meses."),
    ("F5_6meses", "6 meses EC", "183", "Agrega mediciones antropométricas y seguimiento a 6 meses."),
    ("F6_9meses", "9 meses EC", "198", "Agrega el último punto antes del outcome de 12 meses."),
]

outcome_rows = [
    ("Stunting", "Talla baja para la edad", "HAZ < -2 DS", "zscoretalla12cat == 1", "7.623 / 30.953", "24,6%"),
    ("Bajo_peso", "Bajo peso para la edad", "WAZ < -2 DS", "zscorepeso12cat == 1", "3.239 / 29.897", "10,8%"),
    ("Wasting", "Desnutrición aguda", "WHZ < -2 DS", "zscorepesotalla12cat == 1", "1.232 / 29.828", "4,1%"),
]

retention_rows = [
    ("40 semanas", "51.460", "79,4%"),
    ("3 meses", "45.174", "69,7%"),
    ("6 meses", "38.223", "59,0%"),
    ("9 meses", "32.503", "50,2%"),
    ("12 meses", "30.383", "46,9%"),
]

missing_rows = [
    ("Sin nulos", "1", "0,1%", "Mantener."),
    ("Bajo, 1-25%", "261", "34,7%", "Usables con revisión de tipo y limpieza."),
    ("Medio, 26-50%", "181", "24,0%", "Usables si son clínicamente relevantes."),
    ("Alto, 51-75%", "226", "30,0%", "Revisar si faltan por abandono o por protocolo."),
    ("Muy alto, >75%", "84", "11,2%", "Candidatas a exclusión o análisis específico."),
]

coverage_rows = [
    ("Universales", ">=50% de cobertura en todos los periodos", "156", "Base robusta para modelos históricos."),
    ("Parciales", ">=50% en algunos periodos", "457", "Útiles si se modela por cohorte o se toleran faltantes."),
    ("Solo recientes", "Alta cobertura en P5-P6", "403", "Capturan mejoras de registro/protocolo recientes."),
    ("Sin datos", "<10% en todos los periodos", "46", "Bajo valor predictivo salvo justificación clínica fuerte."),
]

prevalence_rows = [
    ("P1 (~1998-2001)", "15,3%", "37,0%", "4,6%"),
    ("P2 (~2002-2004)", "12,9%", "30,4%", "4,7%"),
    ("P3 (~2005-2006)", "12,6%", "34,0%", "3,2%"),
    ("P4 (2007-2012)", "11,0%", "27,1%", "4,0%"),
    ("P5 (2013-2017)", "10,5%", "21,4%", "4,4%"),
    ("P6 (2018-2023)", "9,4%", "19,5%", "3,9%"),
]

correlation_rows = [
    ("zscoretalla0", "-0,270", "-0,199", "-0,078", "Talla muy baja al nacimiento se asocia con mayor stunting."),
    ("zscorepeso0", "-0,242", "-0,200", "-0,100", "Peso bajo respecto al estándar anticipa déficit nutricional."),
    ("RCIUtalla", "+0,237", "+0,178", "+0,069", "Restricción de crecimiento intrauterino en talla."),
    ("RCIUpesoytallanacer", "+0,234", "+0,183", "+0,074", "RCIU combinada de peso y talla."),
    ("RCIUpeso", "+0,205", "+0,167", "+0,079", "RCIU en peso."),
    ("SGAprema", "+0,198", "+0,161", "+0,074", "Pequeño para edad gestacional."),
    ("ERN_Talla", "-0,193", "-0,165", "-0,089", "Longitud al nacer."),
    ("ERN_Peso", "-0,185", "-0,177", "-0,112", "Peso al nacer."),
    ("CP_TallaMadre", "-0,161", "-0,096", "-0,024", "Factor contextual/genético materno."),
]

signal_rows = [
    ("Prenatal/parto", "0,161", "0,037", "7", "41"),
    ("Nacimiento", "0,270", "0,116", "23", "34"),
    ("Hospitalización", "0,104", "0,038", "11", "32"),
    ("40 semanas", "0,263", "0,082", "15", "30"),
    ("3 meses", "0,381", "0,136", "15", "26"),
    ("6 meses", "0,582", "0,165", "10", "19"),
    ("9 meses", "0,641", "0,207", "9", "15"),
]

auc_rows = [
    ("F0 Prenatal/parto", "0,6454", "0,6183", "0,5550"),
    ("F1 Nacimiento", "0,7374", "0,7509", "0,6887"),
    ("F2 Hospitalización", "0,7405", "0,7538", "0,7054"),
    ("F3 40 semanas", "0,7678", "0,7725", "0,7253"),
    ("F4 3 meses", "0,8209", "0,8737", "0,8260"),
    ("F5 6 meses", "0,8935", "0,9360", "0,8950"),
    ("F6 9 meses", "0,9290", "0,9634", "0,9245"),
]

full_metric_rows = [
    ("Stunting", "F0", "0,6454", "0,0102", "0,5088", "0,6891", "0,4133"),
    ("Stunting", "F1", "0,7374", "0,0056", "0,6406", "0,7077", "0,5054"),
    ("Stunting", "F2", "0,7405", "0,0063", "0,6391", "0,7135", "0,5080"),
    ("Stunting", "F3", "0,7678", "0,0081", "0,6214", "0,7583", "0,5262"),
    ("Stunting", "F4", "0,8209", "0,0094", "0,6911", "0,7793", "0,5838"),
    ("Stunting", "F5", "0,8935", "0,0039", "0,7917", "0,8229", "0,6785"),
    ("Stunting", "F6", "0,9290", "0,0031", "0,8342", "0,8594", "0,7368"),
    ("Bajo peso", "F0", "0,6183", "0,0081", "0,3958", "0,7487", "0,2284"),
    ("Bajo peso", "F1", "0,7509", "0,0116", "0,5638", "0,7798", "0,3341"),
    ("Bajo peso", "F2", "0,7538", "0,0115", "0,5495", "0,7959", "0,3404"),
    ("Bajo peso", "F3", "0,7725", "0,0118", "0,5471", "0,8265", "0,3677"),
    ("Bajo peso", "F4", "0,8737", "0,0083", "0,6854", "0,8667", "0,4929"),
    ("Bajo peso", "F5", "0,9360", "0,0040", "0,7894", "0,9123", "0,6287"),
    ("Bajo peso", "F6", "0,9634", "0,0043", "0,8444", "0,9413", "0,7256"),
    ("Wasting", "F0", "0,5550", "0,0086", "0,0771", "0,9431", "0,0467"),
    ("Wasting", "F1", "0,6887", "0,0097", "0,3174", "0,8904", "0,1643"),
    ("Wasting", "F2", "0,7054", "0,0142", "0,3262", "0,8949", "0,1727"),
    ("Wasting", "F3", "0,7253", "0,0079", "0,3311", "0,9036", "0,1858"),
    ("Wasting", "F4", "0,8260", "0,0124", "0,5000", "0,9048", "0,2695"),
    ("Wasting", "F5", "0,8950", "0,0101", "0,6014", "0,9354", "0,3881"),
    ("Wasting", "F6", "0,9245", "0,0100", "0,6835", "0,9488", "0,4766"),
]

test_metric_rows = [
    ("Stunting", "F6 9 meses", "0,9229", "0,8177", "0,8478", "6.191", "1.525"),
    ("Bajo peso", "F6 9 meses", "0,9586", "0,8343", "0,9392", "5.951", "700"),
    ("Wasting", "F6 9 meses", "0,9110", "0,7019", "0,9411", "5.950", "265"),
]

dynamic_risk_rows = [
    ("F0 Prenatal", "0,562", "0,457", "0,106"),
    ("F1 Nacimiento", "0,628", "0,421", "0,207"),
    ("F2 Hospitalización", "0,642", "0,382", "0,260"),
    ("F3 40 semanas", "0,703", "0,338", "0,365"),
    ("F4 3 meses", "0,742", "0,289", "0,454"),
    ("F5 6 meses", "0,767", "0,183", "0,584"),
    ("F6 9 meses", "0,787", "0,145", "0,641"),
]

cohort_prevalence_rows = [
    ("P4 2007-2012", "4.646", "27,1%", "25,9%-28,4%", "4.662", "11,0%", "4.635", "4,0%"),
    ("P5 2013-2017", "9.890", "21,4%", "20,6%-22,2%", "9.892", "10,5%", "9.889", "4,4%"),
    ("P6 2018-2022", "9.226", "19,5%", "18,7%-20,3%", "9.227", "9,4%", "9.226", "3,9%"),
]

cross_cohort_rows = [
    ("P4 2007-2012", "P5 2013-2017", "0,6723", "9.890"),
    ("P4 2007-2012", "P6 2018-2022", "0,6905", "9.226"),
    ("P5 2013-2017", "P4 2007-2012", "0,6667", "4.646"),
    ("P5 2013-2017", "P6 2018-2022", "0,7225", "9.226"),
    ("P6 2018-2022", "P4 2007-2012", "0,6673", "4.646"),
    ("P6 2018-2022", "P5 2013-2017", "0,6903", "9.890"),
]

shap_f6_rows = [
    ("zscoretalla9", "Z-score de talla a 9 meses EC", "1,3241", "La talla reciente domina la predicción de talla baja."),
    ("zscoretalla6", "Z-score de talla a 6 meses EC", "0,6952", "Confirma que el patrón longitudinal importa."),
    ("zscorepeso9", "Z-score de peso a 9 meses EC", "0,1527", "Complementa la señal de talla con estado ponderal."),
    ("velocidad9_6mesesOMS", "Velocidad de crecimiento 6-9 meses", "0,1122", "Detecta desaceleración o recuperación."),
    ("zscoretalla9cat", "Categoría clínica de talla a 9 meses", "0,0928", "Refuerza información de umbrales clínicos."),
    ("zscorepeso6", "Z-score de peso a 6 meses", "0,0580", "Estado de peso intermedio del seguimiento."),
    ("zscoretalla2", "Z-score de talla en 40 semanas", "0,0555", "Señal temprana post-egreso."),
    ("CP_TallaMadre", "Talla materna", "0,0286", "Contexto genético y social no modificable."),
]

dashboard_rows = [
    ("test_predictions.csv", "Predicciones del 20% de test por paciente.", "prob_* por fase, real_* outcomes, variables clínicas de contexto."),
    ("metricas_por_fase.csv", "Métricas de test por fase y outcome.", "AUC_test, Sens_test, Spec_test, n_test, n_pos_test."),
    ("shap_importancia_global.csv", "Ranking global de importancia SHAP.", "feature, mean_abs_shap."),
    ("shap_values.csv", "Valores SHAP individuales por paciente.", "paciente_idx + una columna por feature."),
    ("README.json", "Contrato de lectura para dashboard.", "Fases, outcomes, grupos, descripción de archivos."),
]

resource_rows = [
    ("OMS - Child Growth Standards", "https://www.who.int/tools/child-growth-standards/standards", "Comprender z-scores, curvas y estándares antropométricos."),
    ("LightGBM - Advanced Topics", "https://lightgbm.readthedocs.io/en/latest/Advanced-Topics.html", "Manejo de faltantes y variables categóricas en LightGBM."),
    ("LightGBM - Python API", "https://lightgbm.readthedocs.io/en/v4.0.0/Python-API.html", "Entrenamiento, Booster, Dataset, carga y predicción."),
    ("SHAP - Topical overviews", "https://shap.readthedocs.io/en/latest/overviews.html", "Introducción a interpretabilidad con Shapley values."),
    ("SHAP - Shapley values tutorial", "https://shap.readthedocs.io/en/stable/example_notebooks/overviews/An%20introduction%20to%20explainable%20AI%20with%20Shapley%20values.html", "Tutorial práctico de explicabilidad."),
    ("scikit-learn - ROC AUC", "https://sklearn.org/stable/modules/generated/sklearn.metrics.roc_auc_score.html", "Definición e implementación de ROC-AUC."),
    ("scikit-learn - Precision-Recall", "https://scikit-learn.org/stable/auto_examples/model_selection/plot_precision_recall.html", "Por qué PR es útil con clases desbalanceadas."),
    ("scikit-learn - StratifiedKFold", "https://sklearn.org/stable/modules/generated/sklearn.model_selection.StratifiedKFold.html", "Validación cruzada preservando proporción de clases."),
    ("MLflow Tracking", "https://www.mlflow.org/docs/latest/ml/tracking", "Registro de experimentos, métricas, artefactos y modelos."),
    ("DVC - What is DVC?", "https://dvc.org/doc/user-guide/what-is-dvc", "Versionamiento de datos y modelos en proyectos de ML."),
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color="BFBFBF"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def setup_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10)
    styles["Normal"].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    styles["Normal"].paragraph_format.space_after = Pt(4)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[style_name].font.name = "Calibri"
        styles[style_name].font.color.rgb = RGBColor(31, 78, 121)
        styles[style_name].paragraph_format.keep_with_next = True
    styles["Heading 1"].font.size = Pt(15)
    styles["Heading 2"].font.size = Pt(12.5)
    styles["Heading 3"].font.size = Pt(11)
    styles["List Bullet"].font.name = "Calibri"
    styles["List Bullet"].font.size = Pt(10)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(8 if level == 1 else 5)
    p.paragraph_format.space_after = Pt(3)
    return p


def add_para(doc, text, style=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.add_run(item)


def add_code(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(code)
    r.font.name = "Consolas"
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(50, 50, 50)
    p.paragraph_format.left_indent = Inches(0.25)
    return p


def style_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for r_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_border(cell)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                for run in p.runs:
                    run.font.size = Pt(8.5)
            if r_idx == 0:
                set_cell_shading(cell, "D9EAF7")
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(31, 78, 121)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    style_table(table)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    return table


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(89, 89, 89)


def add_image(doc, key, caption, width=6.4):
    path = IMAGES[key]
    if not path.exists():
        add_para(doc, f"[Figura no encontrada: {path}]")
        return
    doc.add_picture(str(path), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_caption(doc, caption)


def add_title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Documentación técnica completa del proyecto")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(31, 78, 121)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Predicción de malnutrición a 12 meses de edad corregida en niños prematuros")
    r.bold = True
    r.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("PMCI / Fundación Canguro · Versión ampliada para transferencia técnica · Abril de 2026")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    p.add_run("Audiencia objetivo: programador con nociones básicas de inteligencia artificial")


def add_footer(doc):
    footer = doc.sections[0].footer.paragraphs[0]
    footer.text = "Documentación técnica completa - PMCI / Fundación Canguro"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(127, 127, 127)


def build_document():
    doc = Document()
    setup_styles(doc)
    add_title(doc)
    add_footer(doc)

    add_heading(doc, "0. Cómo leer este documento", 1)
    add_para(doc, "Este documento está escrito para que una persona que no participó en el proyecto pueda reconstruir el razonamiento técnico completo: qué problema clínico se aborda, qué datos existen, cómo se limpiaron, cómo se construyeron los modelos, cómo se evaluaron, qué resultados se obtuvieron, qué riesgos hay y cómo continuar.")
    add_para(doc, "La documentación mezcla tres niveles: contexto clínico, decisiones de ingeniería de datos y fundamentos de inteligencia artificial. Cuando aparece un concepto técnico, se explica primero con lenguaje práctico y luego se conecta con su uso concreto en este proyecto.")
    add_bullets(doc, [
        "Si se quiere entender el propósito, leer secciones 1 a 3.",
        "Si se quiere ejecutar o mantener el pipeline, leer secciones 4 a 12.",
        "Si se quiere interpretar resultados y defender decisiones, leer secciones 13 a 18.",
        "Si se quiere seguir aprendiendo, usar las referencias y recursos de la sección 21.",
    ])

    add_heading(doc, "1. Contexto clínico y motivación", 1)
    add_para(doc, "El proyecto se desarrolla con datos del Programa Madre Canguro Integral (PMCI) de la Fundación Canguro. El PMCI atiende y sigue niños prematuros o con bajo peso al nacer, una población con riesgo aumentado de complicaciones neonatales, dificultades de crecimiento y problemas de desarrollo. El Método Madre Canguro combina contacto piel a piel, lactancia materna y seguimiento ambulatorio de alto riesgo.")
    add_para(doc, "El problema específico es anticipar malnutrición a los 12 meses de edad corregida. Edad corregida significa que, para un niño prematuro, la edad se ajusta tomando como referencia las 40 semanas de gestación. Por ejemplo, un niño nacido 8 semanas antes de término tendrá 4 meses cronológicos pero 2 meses corregidos. En seguimiento de prematuros, esta corrección es esencial para no comparar su crecimiento con una edad biológica que no corresponde.")
    add_para(doc, "La motivación clínica es intervenir antes. Si el riesgo se identifica hasta el control de 12 meses, la herramienta llega tarde. Si el riesgo puede estimarse desde nacimiento, 40 semanas, 3, 6 o 9 meses, el equipo clínico puede priorizar controles, nutrición, educación familiar o seguimiento intensivo.")

    add_heading(doc, "2. Objetivo general y entregables del proyecto", 1)
    add_para(doc, "El objetivo general es identificar factores clínicos, perinatales y sociodemográficos asociados a malnutrición en niños prematuros evaluados a 12 meses de edad corregida, usando modelos de aprendizaje automático sobre datos longitudinales del PMCI.")
    add_para(doc, "En la práctica, el proyecto ya produjo cuatro tipos de entregables: análisis exploratorio de datos, modelos predictivos por fase temporal, artefactos de interpretabilidad y archivos exportables para alimentar una futura plataforma o dashboard.")
    add_table(doc, ["Entregable", "Dónde está", "Para qué sirve"], [
        ("EDA de prematurez", "notebooks/EDA_Prematurez.ipynb", "Entender distribución de prematuridad, faltantes y variables clave iniciales."),
        ("EDA de malnutrición", "notebooks/EDA_Malnutricion_ejecutado.ipynb", "Estudiar outcomes, deserción, cobertura histórica, prevalencias y correlaciones."),
        ("Pipeline de modelado", "notebooks/Modelado_Malnutricion-final.ipynb", "Entrenar, evaluar, interpretar y guardar modelos."),
        ("Modelos finales", "modelos_pmci/modelo_*.lgb", "Modelos LightGBM listos para inferencia."),
        ("Datos para dashboard", "dashboard_data/*.csv y README.json", "Predicciones, métricas, SHAP y contrato de lectura para interfaz."),
        ("Documentación técnica", "documentacion/HANDOVER.md y pipeline_documentacion.md", "Resumen de decisiones, ejecución y próximos pasos."),
    ], widths=[1.5, 2.1, 3.0])

    add_heading(doc, "3. Repositorio y mapa de archivos", 1)
    add_para(doc, "La estructura del repositorio separa datos originales, notebooks, modelos entrenados, datos para dashboard y documentos de transferencia. Esta separación es importante porque permite distinguir entre fuente de verdad, experimentación y artefactos listos para consumo.")
    add_code(doc, """proyecto_desarrollo_soluciones_canguro/
├── data/
│   ├── KMC-70k-93-2024-Malnutricion-conVel-DATA-SPSS-20250322.xlsx
│   ├── KMC-70K-diccionarioVARS-Malnutricion-PhETI-rev20250520-MAIA.xlsx
│   └── feature_plan.json
├── notebooks/
│   ├── EDA_Prematurez.ipynb
│   ├── EDA_Malnutricion_ejecutado.ipynb
│   ├── Mejoras_Modelo.ipynb
│   └── Modelado_Malnutricion-final.ipynb
├── modelos_pmci/
│   ├── modelo_{outcome}_{fase}.lgb
│   ├── features_{fase}.json
│   └── metadata.json
├── dashboard_data/
│   ├── test_predictions.csv
│   ├── metricas_por_fase.csv
│   ├── shap_importancia_global.csv
│   ├── shap_values.csv
│   └── README.json
└── documentacion/
    ├── HANDOVER.md
    ├── pipeline_documentacion.md
    └── propuestas PDF""")
    add_para(doc, "Un programador que reciba este proyecto debería comenzar leyendo documentacion/HANDOVER.md y dashboard_data/README.json. El primero explica el estado del trabajo; el segundo define el contrato de archivos que podría consumir una aplicación.")

    add_heading(doc, "4. Fuente de datos y diccionario", 1)
    add_para(doc, "La fuente principal es un Excel exportado desde SPSS con 64.801 filas y 753 columnas. Cada fila representa una historia clínica o registro de paciente. Las columnas cubren identificación anonimizada, condiciones sociales, control prenatal, parto, nacimiento, hospitalización y seguimiento ambulatorio.")
    add_para(doc, "El diccionario PhETI documenta 640 variables en la hoja VARS-(KMC70k). No todas las variables del Excel principal están completamente documentadas, y no todas las variables documentadas son necesariamente aptas para modelado. Por eso se combinaron tres fuentes: nombres de columnas reales, diccionario clínico y revisión técnica de leakage.")
    add_table(doc, ["Hoja del diccionario", "Contenido", "Uso dentro del proyecto"], [
        ("PhETI-KMC-Dictionary-ReadMe", "Notas generales del diccionario.", "Entender la convención del archivo."),
        ("VARS-(KMC70k)", "Variables, descripción, tipo, fase, unidades y faltantes.", "Fuente principal para comprender variables."),
        ("Phases", "Fases del proceso clínico.", "Relacionar variables con momentos temporales."),
        ("Episodes", "Episodios clínicos.", "Agrupar eventos del PMCI."),
        ("TopicsOfInterest", "Temas de interés clínico.", "Explorar grupos conceptuales de variables."),
        ("PREFIX-VARIABLES", "Prefijos como CP, CSP, ERN, HD.", "Leer rápidamente el origen de una variable."),
    ], widths=[2.0, 2.2, 2.5])
    add_table(doc, ["Prefijo", "Significado aproximado"], [
        ("IDEN", "Identificación del sujeto."),
        ("CP", "Control prenatal."),
        ("CSP", "Condición social."),
        ("ERN", "Recién nacido."),
        ("HD", "Hospitalización."),
        ("PA", "Parto."),
        ("PCANG", "Posición canguro."),
        ("NUT", "Nutrición."),
        ("CR", "Crecimiento."),
        ("OX", "Oxígeno."),
    ], widths=[1.0, 5.0])
    add_para(doc, "El diccionario clasifica variables en tipos primarios: categóricas, booleanas, flotantes, enteras, texto y fechas. Esto afecta el preprocesamiento: una variable numérica puede convertirse directamente; una categórica requiere codificación; una fecha normalmente no se usa como número bruto salvo que se derive una característica temporal con sentido clínico.")

    add_heading(doc, "5. Definición clínica de los desenlaces", 1)
    add_para(doc, "El proyecto usa tres desenlaces binarios de malnutrición basados en z-scores. Un z-score indica cuántas desviaciones estándar está una medición por encima o por debajo del valor esperado para edad y sexo. Un valor menor a -2 DS suele interpretarse como déficit clínicamente relevante.")
    add_para(doc, "Para un programador: el z-score convierte medidas como peso o talla en una escala comparable. Dos niños pueden tener el mismo peso, pero si tienen edades corregidas distintas, el significado clínico cambia. Por eso el modelo no usa solo gramos o centímetros; también usa variables derivadas que comparan al niño contra estándares de crecimiento.")
    add_table(doc, ["Outcome", "Significado", "Criterio", "Variable base", "Positivos", "Prevalencia"], outcome_rows, widths=[1.0, 1.5, 1.0, 1.5, 1.0, 0.8])
    add_para(doc, "Stunting fue tratado como outcome principal porque es frecuente, clínicamente importante y muestra una señal predictiva consistente. Bajo peso también alcanzó alto desempeño. Wasting es más escaso, por lo que se vuelve más difícil entrenar modelos sensibles sin aumentar falsas alarmas.")

    add_heading(doc, "6. Preparación de datos", 1)
    add_para(doc, "La preparación de datos tuvo cuatro objetivos: convertir valores faltantes, asegurar tipos correctos, construir outcomes y evitar fuga de información. La fuga de información ocurre cuando el modelo recibe variables que no estarían disponibles al momento real de predicción. En salud esto puede producir métricas artificialmente altas y una herramienta inútil en práctica.")
    add_bullets(doc, [
        "Los valores '#NULL!' provenientes de SPSS se convirtieron a NaN.",
        "Las columnas con valores numéricos escritos como texto se convirtieron a numéricas cuando era razonable.",
        "Los outcomes se construyeron desde las categorías de z-score a 12 meses.",
        "Se excluyeron variables medidas a 12 meses o posteriores a la fase evaluada.",
        "Para los modelos finales se permitió que LightGBM manejara faltantes de forma nativa.",
    ])
    add_table(doc, ["Nivel de faltantes", "Columnas", "% del total", "Decisión práctica"], missing_rows, widths=[1.5, 0.8, 0.8, 3.5])
    add_image(doc, "missing_values", "Figura 1. Distribución de valores faltantes observada en el EDA de prematurez.", width=6.2)
    add_para(doc, "Una decisión importante fue no imputar todo de manera agresiva en el pipeline final. Imputar significa reemplazar datos faltantes por valores estimados, como mediana o moda. Esto puede ayudar a modelos que no aceptan NaN, pero también puede introducir señales artificiales. LightGBM puede aprender rutas específicas para valores faltantes, lo que reduce la necesidad de inventar datos.")

    add_heading(doc, "7. EDA de prematurez", 1)
    add_para(doc, "El EDA de prematurez no es el modelo final de malnutrición, pero fue útil para entender la base. Mostró que la variable preterm está casi completa: solo 16 nulos en 64.801 registros. La distribución fue 48.320 prematuros (74,6%) y 16.465 no prematuros (25,4%), con una relación aproximada 3:1.")
    add_image(doc, "preterm_distribution", "Figura 2. Distribución de la variable preterm.", width=6.0)
    add_para(doc, "El desbalance de clases significa que una clase aparece mucho más que otra. Si un modelo se entrena sin cuidado, puede aprender a favorecer la clase mayoritaria. En este caso, para prematurez la mayoría son prematuros; en malnutrición ocurre lo contrario para algunos outcomes, como wasting, donde los positivos son pocos.")
    add_image(doc, "prematurity_key_vars", "Figura 3. Variables clave relacionadas con prematurez.", width=6.4)
    add_image(doc, "missing_strategy", "Figura 4. Estrategia inicial propuesta para manejo de nulos.", width=6.2)
    add_image(doc, "numeric_distributions", "Figura 5. Distribución de variables numéricas clave por clase.", width=6.4)
    add_para(doc, "La principal lección del EDA de prematurez fue que la base tiene suficiente tamaño para modelado, pero exige cuidado con tipos, faltantes y clases desbalanceadas. También mostró variables iniciales con relevancia clínica, como edad gestacional, peso al nacer, talla, SGAprema y edad materna.")

    add_heading(doc, "8. EDA de malnutrición", 1)
    add_para(doc, "El EDA de malnutrición estudió directamente los desenlaces de 12 meses. La primera observación crítica es que aproximadamente la mitad de los pacientes no tiene outcome completo a 12 meses. Esto no debe interpretarse como simple ruido: en seguimiento clínico real, la ausencia de dato puede estar asociada a deserción, dificultades de acceso, cambios de protocolo o características del paciente.")
    add_table(doc, ["Visita", "Pacientes que llegaron", "% del total"], retention_rows, widths=[1.6, 1.4, 1.0])
    add_para(doc, "La retención cae de 79,4% en 40 semanas a 46,9% en 12 meses. Esto afecta la interpretación de prevalencias y modelos: los pacientes con outcome disponible podrían no representar perfectamente a todos los niños iniciales. En términos estadísticos, es necesario preguntarse si los faltantes son MCAR, MAR o MNAR.")
    add_bullets(doc, [
        "MCAR: faltan completamente al azar; raro en salud.",
        "MAR: faltan por variables observables, por ejemplo sede, año, edad materna o distancia.",
        "MNAR: faltan por razones relacionadas con el propio desenlace, por ejemplo familias de niños más complicados o más sanos dejan de asistir.",
    ])
    add_table(doc, ["Categoría de cobertura", "Regla", "Variables", "Implicación"], coverage_rows, widths=[1.3, 2.1, 0.8, 2.8])
    add_para(doc, "La cobertura también cambió por periodos históricos. Se identificaron 156 variables universales con cobertura razonable en todos los periodos, pero 467 variables tuvieron cambios de cobertura superiores a 30 puntos porcentuales. Esto sugiere cambios de protocolo, cambios de registro o disponibilidad progresiva de variables.")
    add_table(doc, ["Periodo", "Bajo peso", "Stunting", "Wasting"], prevalence_rows, widths=[1.8, 1.0, 1.0, 1.0])
    add_para(doc, "La prevalencia de stunting disminuye de forma importante en el tiempo: 37,0% en P1, 27,1% en P4 y 19,5% en P6. Bajo peso también disminuye. Wasting se mantiene más estable alrededor de 3-5%. Esta tendencia puede reflejar mejoras del PMCI, cambios poblacionales o cambios de captura de datos.")
    add_table(doc, ["Variable", "r stunting", "r bajo peso", "r wasting", "Lectura"], correlation_rows, widths=[1.5, 0.8, 0.8, 0.8, 2.7])
    add_para(doc, "Las correlaciones no son causalidad. Una correlación negativa entre zscoretalla0 y stunting significa que valores más altos de talla al nacimiento se asocian con menor probabilidad de talla baja a 12 meses. No significa que aumentar artificialmente esa variable en una tabla cambiaría el resultado. Para causalidad se necesitarían diseños y supuestos adicionales.")
    add_table(doc, ["Fase", "Máx |r|", "Media |r|", "Vars >0,05", "Vars totales"], signal_rows, widths=[1.5, 0.8, 0.8, 0.9, 0.9])
    add_para(doc, "La señal predictiva crece a medida que avanza el seguimiento, especialmente a 6 y 9 meses. Esto anticipa el resultado principal del modelado: mientras más cerca del outcome y más mediciones de crecimiento hay, mejor predice el modelo.")

    add_heading(doc, "9. Plan de features y cascada temporal", 1)
    add_para(doc, "El proyecto usa una cascada temporal. La idea es simple: entrenar un modelo por momento clínico usando solo la información acumulada hasta ese momento. Esto evita mezclar información futura y permite que la herramienta sea útil en distintas etapas del PMCI.")
    add_table(doc, ["Fase", "Momento", "N features", "Contenido principal"], phase_rows, widths=[1.5, 1.4, 0.8, 3.2])
    add_para(doc, "La cascada también ayuda a responder una pregunta de investigación: ¿en qué momento emerge la señal predictiva? Si el desempeño fuera alto desde F1, el nacimiento sería suficiente para riesgo temprano. Si solo mejora en F6, entonces el seguimiento longitudinal es indispensable.")

    add_heading(doc, "10. Conceptos de IA necesarios para entender el pipeline", 1)
    add_heading(doc, "10.1 Clasificación supervisada", 2)
    add_para(doc, "El proyecto es un problema de clasificación supervisada. Supervisada significa que cada ejemplo de entrenamiento tiene una etiqueta conocida: 0 o 1 para cada outcome. Clasificación significa que el modelo debe asignar una clase o una probabilidad de pertenecer a la clase de riesgo.")
    add_heading(doc, "10.2 Probabilidad, umbral y clase", 2)
    add_para(doc, "El modelo produce una probabilidad, por ejemplo P(stunting)=0,72. Para convertirla en alerta se necesita un umbral, por ejemplo 0,50. Si la probabilidad supera el umbral, se etiqueta como riesgo. En salud el umbral no debe elegirse solo por conveniencia: depende de cuánto cuesta pasar por alto un caso real frente a generar una alerta innecesaria.")
    add_heading(doc, "10.3 Validación cruzada", 2)
    add_para(doc, "La validación cruzada divide el conjunto de entrenamiento en varias partes. En cada iteración se entrena con unas partes y se valida con otra. En este proyecto se usaron 5 folds estratificados, lo que preserva aproximadamente la proporción de positivos y negativos en cada fold. Esto es importante porque los outcomes están desbalanceados.")
    add_heading(doc, "10.4 LightGBM", 2)
    add_para(doc, "LightGBM es un algoritmo de gradient boosting basado en árboles de decisión. Un árbol divide los datos con preguntas del tipo '¿zscoretalla9 < -1,8?'. Gradient boosting construye muchos árboles pequeños de forma secuencial: cada nuevo árbol intenta corregir errores de los anteriores. LightGBM está optimizado para datos tabulares grandes y maneja NaN de forma nativa.")
    add_para(doc, "En el proyecto se eligió LightGBM porque la base es tabular, tiene muchas variables, contiene faltantes y puede incluir relaciones no lineales entre crecimiento, edad gestacional, alimentación y hospitalización.")
    add_heading(doc, "10.5 SHAP", 2)
    add_para(doc, "SHAP es una técnica de explicabilidad. Asigna a cada variable una contribución aproximada sobre la predicción de un paciente o sobre la importancia global del modelo. La intuición viene de teoría de juegos: si las variables fueran jugadores que colaboran para producir una predicción, SHAP estima cuánto aportó cada jugador.")
    add_para(doc, "En este proyecto, SHAP sirve para revisar si el modelo está aprendiendo patrones clínicamente razonables. Por ejemplo, si las variables más importantes para stunting son talla a 9 meses, talla a 6 meses y velocidad de crecimiento, la explicación es coherente con el problema.")

    add_heading(doc, "11. Construcción del pipeline final", 1)
    add_para(doc, "El notebook Modelado_Malnutricion-final.ipynb implementa el flujo principal. El pipeline carga datos, construye outcomes, arma features acumuladas, entrena modelos LightGBM por fase y outcome, evalúa resultados, calcula SHAP, analiza riesgo dinámico, compara baseline, analiza cohortes y guarda modelos finales.")
    add_code(doc, """Flujo lógico:
1. Cargar Excel principal y feature_plan.json.
2. Reemplazar '#NULL!' por NaN.
3. Construir stunting12m, underweight12m_b y wasting12m.
4. Excluir variables con leakage.
5. Definir cumulative_features por F0-F6.
6. Para cada outcome y fase:
   a. Filtrar pacientes con outcome disponible.
   b. Entrenar LightGBM con StratifiedKFold.
   c. Guardar métricas y predicciones out-of-fold.
7. Comparar AUC, sensibilidad, especificidad y F1.
8. Calcular SHAP.
9. Entrenar modelos finales y guardarlos en modelos_pmci/.
10. Exportar dashboard_data/.""")
    add_para(doc, "Un detalle clave es que se entrenan modelos independientes. No hay un único modelo que recibe una fase como parámetro. Esto simplifica la inferencia: si el paciente está en F1, se carga el modelo F1; si está en F6, se carga el modelo F6.")

    add_heading(doc, "12. Métricas de evaluación", 1)
    add_para(doc, "Las métricas se eligieron para clasificación clínica desbalanceada. Ninguna métrica cuenta toda la historia. Por eso el documento reporta AUC, sensibilidad, especificidad y F1.")
    add_table(doc, ["Métrica", "Qué responde", "Interpretación práctica"], [
        ("AUC ROC", "¿El modelo ordena mejor a pacientes positivos que negativos?", "0,5 es azar; 1,0 es separación perfecta. Útil para comparar modelos sin fijar umbral."),
        ("Sensibilidad", "De los casos reales, ¿cuántos detecta?", "Alta sensibilidad reduce falsos negativos."),
        ("Especificidad", "De los no casos, ¿cuántos descarta?", "Alta especificidad reduce falsas alarmas."),
        ("F1", "Balance entre precisión y sensibilidad.", "Útil con desbalance, pero depende del umbral."),
        ("Curva PR", "Relación precisión-sensibilidad.", "Más informativa que ROC cuando los positivos son pocos."),
    ], widths=[1.2, 2.2, 3.2])
    add_para(doc, "En un contexto médico, la sensibilidad suele ser prioritaria para tamizaje porque un falso negativo puede implicar no intervenir a un niño en riesgo. Sin embargo, demasiados falsos positivos pueden saturar al equipo clínico. La selección final de umbral debe hacerse con expertos.")

    add_heading(doc, "13. Resultados del modelado", 1)
    add_para(doc, "Los resultados validan la arquitectura temporal. El AUC aumenta con cada fase. En F0, la predicción es moderada; en F5 y F6 se vuelve alta. Esto confirma que las mediciones longitudinales de crecimiento contienen mucha información sobre el estado nutricional a 12 meses.")
    add_table(doc, ["Fase", "Stunting", "Bajo peso", "Wasting"], auc_rows, widths=[1.8, 1.0, 1.0, 1.0])
    add_image(doc, "auc_cascade", "Figura 6. Evolución del AUC por fase y outcome.", width=6.4)
    add_table(doc, ["Outcome", "Fase", "AUC", "Std", "Sens", "Spec", "F1"], full_metric_rows, widths=[1.0, 0.5, 0.7, 0.6, 0.7, 0.7, 0.7])
    add_para(doc, "La tabla completa muestra tres patrones. Primero, stunting tiene desempeño útil desde nacimiento, pero mejora mucho con 6 y 9 meses. Segundo, bajo peso es el outcome con mayor AUC final. Tercero, wasting tiene AUC final alto pero sensibilidad menor, lo cual se explica por su baja prevalencia.")
    add_image(doc, "stunting_metrics", "Figura 7. AUC, sensibilidad y especificidad para stunting.", width=6.2)
    add_image(doc, "roc_curves", "Figura 8. Curvas ROC por fase para stunting.", width=5.6)
    add_image(doc, "pr_curves", "Figura 9. Curvas precision-recall por outcome.", width=6.4)
    add_table(doc, ["Outcome", "Mejor fase", "AUC test", "Sens.", "Spec.", "n test", "positivos"], test_metric_rows, widths=[1.2, 1.0, 0.8, 0.7, 0.7, 0.8, 0.8])
    add_para(doc, "El split de test confirma el desempeño observado en validación cruzada, aunque las cifras no son idénticas porque se trata de una evaluación 80/20 específica. Esto es normal: validación cruzada estima desempeño promedio; test evalúa una partición final reservada.")

    add_heading(doc, "14. Interpretabilidad con SHAP", 1)
    add_para(doc, "La interpretabilidad es crucial porque una herramienta clínica no debería limitarse a decir 'riesgo alto'. Debe permitir revisar qué factores empujan la predicción hacia riesgo y si esos factores son coherentes. En este proyecto, SHAP se calculó para stunting en F6 usando una muestra de 3.000 pacientes y 198 features.")
    add_image(doc, "shap_summary", "Figura 10. SHAP summary plot: dirección e importancia de variables.", width=5.8)
    add_image(doc, "shap_bar", "Figura 11. Importancia global SHAP para stunting en F6.", width=5.8)
    add_table(doc, ["Variable", "Descripción", "SHAP medio", "Lectura"], shap_f6_rows, widths=[1.5, 2.0, 0.8, 2.7])
    add_para(doc, "La lectura principal es que el modelo no parece depender de variables arbitrarias: usa sobre todo la trayectoria de crecimiento del niño. Esto es clínicamente razonable. La ventana de 6 a 9 meses aparece como un momento crítico para intervención antes del desenlace de 12 meses.")
    add_para(doc, "Advertencia importante: SHAP explica el modelo, no demuestra causalidad. Si una variable tiene alto SHAP, significa que el modelo la usa para predecir; no significa que modificar esa variable en la realidad cambie necesariamente el desenlace.")

    add_heading(doc, "15. Sistema de riesgo dinámico", 1)
    add_para(doc, "El sistema de riesgo dinámico calcula probabilidades de stunting para el mismo paciente en varias fases. Esto permite observar trayectorias: algunos pacientes mantienen riesgo alto, otros mejoran, y otros empeoran con el tiempo.")
    add_image(doc, "dynamic_risk", "Figura 12. Trayectorias individuales y distribución de riesgo por fase.", width=6.4)
    add_table(doc, ["Fase", "P media stunted", "P media normal", "Delta"], dynamic_risk_rows, widths=[1.5, 1.2, 1.2, 0.8])
    add_para(doc, "El delta entre grupos aumenta con cada fase: de 0,106 en F0 a 0,641 en F6. Esto quiere decir que, conforme se agregan mediciones, las probabilidades promedio de niños que terminarán con stunting se separan cada vez más de las de niños normales.")

    add_heading(doc, "16. Comparación con baseline", 1)
    add_para(doc, "Se comparó LightGBM contra regresión logística L1 para stunting en F6. La regresión logística L1 es un baseline útil porque es más interpretable y selecciona variables al empujar algunos coeficientes a cero.")
    add_table(doc, ["Modelo", "AUC CV", "Comentario"], [
        ("Regresión logística L1", "0,9216 ± 0,0031", "Modelo lineal, más simple e interpretable."),
        ("LightGBM", "0,9290", "Mejor desempeño, maneja faltantes y no linealidades."),
        ("Ganancia", "+0,0074", "Mejora real pero moderada."),
    ], widths=[1.8, 1.4, 3.4])
    add_para(doc, "La ganancia de LightGBM es moderada. Esto no es negativo: sugiere que gran parte de la señal es estable y capturable por modelos más simples. Sin embargo, LightGBM sigue siendo conveniente para producción por su manejo de NaN y su flexibilidad.")

    add_heading(doc, "17. Cohortes temporales y drift", 1)
    add_para(doc, "El análisis de cohortes temporales revisa si la población y el desempeño cambian con los años. Esto es esencial porque un modelo entrenado en registros antiguos puede degradarse si los protocolos, los pacientes o la forma de registrar datos cambian.")
    add_table(doc, ["Cohorte", "n stunting", "Prev. stunting", "IC95 stunting", "n bajo peso", "Prev. bajo peso", "n wasting", "Prev. wasting"], cohort_prevalence_rows, widths=[1.2, 0.8, 0.9, 1.2, 0.8, 0.9, 0.8, 0.9])
    add_image(doc, "cohort_analysis", "Figura 13. AUC por cohorte, prevalencias y comparación temporal.", width=6.4)
    add_para(doc, "La disminución de prevalencia de stunting entre P4 y P6 puede indicar mejora clínica, pero también puede reflejar diferencias de registro. Por eso, el documento no lo interpreta como causalidad del programa sin validación adicional.")
    add_table(doc, ["Train", "Test", "AUC", "n test"], cross_cohort_rows, widths=[1.6, 1.6, 0.8, 0.8])
    add_image(doc, "cross_cohort", "Figura 14. Validación cruzada entre cohortes para stunting en F2.", width=4.8)
    add_image(doc, "feature_drift", "Figura 15. Features con mayor cambio de cobertura entre cohortes.", width=6.2)
    add_para(doc, "La validación cruzada entre cohortes produce AUC entre 0,667 y 0,723 para F2. Estos valores son más bajos que la validación aleatoria del conjunto completo, lo cual es esperable: probar en otro periodo es más difícil y revela drift temporal.")

    add_heading(doc, "18. Inferencia y uso de modelos guardados", 1)
    add_para(doc, "Los modelos finales se guardaron en modelos_pmci/ como archivos .lgb, el formato nativo de LightGBM. Cada fase tiene además un archivo JSON con la lista exacta de features esperadas. Esto es importante: en inferencia, el DataFrame debe tener las columnas en el mismo orden que el entrenamiento.")
    add_code(doc, """import lightgbm as lgb
import json
import pandas as pd
import numpy as np

def predecir_riesgo(paciente, fase, outcome="Stunting", models_dir="modelos_pmci", umbral=0.5):
    model = lgb.Booster(model_file=f"{models_dir}/modelo_{outcome}_{fase}.lgb")
    features = json.load(open(f"{models_dir}/features_{fase}.json", encoding="utf-8"))["features"]
    X = pd.DataFrame([paciente]).reindex(columns=features, fill_value=np.nan)
    prob = float(model.predict(X)[0])
    return {"probabilidad": prob, "riesgo": prob >= umbral}""")
    add_image(doc, "inference_demo", "Figura 16. Ejemplo de cascada de riesgo para un paciente.", width=6.0)
    add_para(doc, "Si un paciente no tiene todas las variables, se completan con NaN. LightGBM puede manejar esos NaN. Si el diccionario del paciente trae variables nuevas no vistas en entrenamiento, se ignoran al reindexar. Esto hace que la función sea flexible para una interfaz clínica.")

    add_heading(doc, "19. Datos exportados para dashboard", 1)
    add_para(doc, "La carpeta dashboard_data/ es el puente entre modelado y aplicación. Contiene archivos planos que una interfaz puede cargar sin depender de notebooks.")
    add_table(doc, ["Archivo", "Qué contiene", "Columnas clave"], dashboard_rows, widths=[1.8, 2.3, 2.5])
    add_para(doc, "test_predictions.csv tiene una fila por paciente del test y columnas de probabilidad para cada outcome y fase. Esto permite construir vistas como 'riesgo por fase', 'riesgo actual', 'comparación con outcome real' o 'trayectoria individual'.")

    add_heading(doc, "20. Experimentos previos y MLflow", 1)
    add_para(doc, "Además del pipeline final, el repositorio contiene rastros de experimentos previos en mlruns/. Esos runs entrenaron modelos como Gradient Boosting sobre una versión más temprana del problema, con target indexnutricion12meses, 10.000 filas, 171 features y métricas como AUC 0,8141 y F1 0,8329.")
    add_para(doc, "Estos experimentos son útiles como historia del proyecto, pero el pipeline final los supera conceptualmente porque usa outcomes clínicos específicos (stunting, bajo peso, wasting), cascada temporal por fases, control explícito de leakage, SHAP y modelos guardados por fase.")
    add_para(doc, "MLflow sirve para registrar experimentos: parámetros, métricas, artefactos y modelos. Si el proyecto continúa, conviene que los notebooks finales también registren cada corrida en MLflow para poder comparar versiones, especialmente si cambian features, cohortes, umbrales o hiperparámetros.")

    add_heading(doc, "21. Limitaciones, riesgos y buenas prácticas", 1)
    add_bullets(doc, [
        "Faltantes y deserción: el 53,1% no llega a 12 meses; esto puede sesgar prevalencias y entrenamiento.",
        "Cercanía temporal: F6 predice muy bien, pero está cerca del outcome; no reemplaza una alerta neonatal temprana.",
        "Desbalance: wasting tiene pocos positivos; sensibilidad y umbral requieren discusión clínica.",
        "Leakage: se excluyeron variables de 12 meses, pero cualquier nueva feature derivada debe auditarse.",
        "Drift temporal: cambios de protocolo y cobertura pueden afectar desempeño futuro.",
        "Interpretabilidad: SHAP no es causalidad; se debe evitar vender factores predictivos como causas.",
        "Ética: el modelo debe apoyar decisiones, no automatizar diagnósticos sin revisión médica.",
    ])
    add_para(doc, "Antes de producción se recomienda calibración de probabilidades, validación temporal más estricta, evaluación por sede, auditoría de sesgos, revisión clínica de variables y pruebas de usabilidad con neonatólogos.")

    add_heading(doc, "22. Recursos recomendados para profundizar", 1)
    add_para(doc, "Los siguientes recursos son útiles para una persona que quiera profundizar en los conceptos técnicos y clínicos usados en el proyecto. Se priorizan fuentes oficiales o documentación primaria.")
    add_table(doc, ["Recurso", "URL", "Para qué usarlo"], resource_rows, widths=[1.7, 2.8, 2.2])

    add_heading(doc, "23. Referencias internas y bibliográficas", 1)
    refs = [
        "Fundación Canguro. Propuesta de proyecto: Identificación de factores de riesgo de malnutrición en niños prematuros evaluados a 12 meses de edad corregida. Documento interno.",
        "Equipo del proyecto. HANDOVER.md. Predicción de malnutrición a 12 meses EC, PMCI/Fundación Canguro. Abril de 2026.",
        "Equipo del proyecto. pipeline_documentacion.md. Pipeline de modelado para predicción de malnutrición a 12 meses EC. Abril de 2026.",
        "Charpak, N., Ruiz, J. G., Zupan, J., et al. (2005). Kangaroo mother care: 25 years after. Acta Paediatrica, 94(5), 514-522.",
        "Charpak, N., Tessier, R., Ruiz, J. G., et al. (2017). Twenty-year follow-up of kangaroo mother care versus traditional care. Pediatrics, 139(1), e20162063.",
        "Charpak, N., & Montealegre-Pomar, A. (2023). Follow-up of Kangaroo Mother Care programmes in the last 28 years: results from a cohort of 57,154 low-birth-weight infants in Colombia. BMJ Global Health, 8(5).",
        "Fenton, T. R., & Kim, J. H. (2013). A systematic review and meta-analysis to revise the Fenton growth chart for preterm infants. BMC Pediatrics, 13, 59.",
        "World Health Organization. (2006). WHO Child Growth Standards.",
        "LightGBM documentation. Advanced Topics and Python API.",
        "SHAP documentation. Explainable AI with Shapley values.",
        "scikit-learn documentation. ROC-AUC, Precision-Recall and StratifiedKFold.",
        "MLflow documentation. Tracking experiments, metrics, artifacts and models.",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        r = p.add_run(ref)
        r.font.size = Pt(9)

    doc.save(OUT)
    return doc


if __name__ == "__main__":
    document = build_document()
    print(OUT)
    print(f"Documento generado con {len(document.paragraphs)} párrafos, {len(document.tables)} tablas.")
