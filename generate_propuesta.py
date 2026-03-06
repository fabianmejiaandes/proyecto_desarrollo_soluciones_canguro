"""
Script to generate the Propuesta Metodologica Word document for Fundacion Canguro.
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os


def set_cell_shading(cell, color_hex):
    """Set background color for a table cell."""
    shading_elm = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    """Set borders for a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')

    for border_name, border_val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if border_val:
            border_elm = parse_xml(
                f'<w:{border_name} {nsdecls("w")} w:val="single" w:sz="4" w:space="0" w:color="{border_val}"/>'
            )
            tcBorders.append(border_elm)

    tcPr.append(tcBorders)


def format_table(table, header_color="2F5496", alt_color="D6E4F0"):
    """Format a table with header styling and alternating row colors."""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Style header row
    for cell in table.rows[0].cells:
        set_cell_shading(cell, header_color)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10)
                run.font.name = "Calibri"

    # Style data rows with alternating colors
    for i, row in enumerate(table.rows[1:], start=1):
        bg = alt_color if i % 2 == 1 else "FFFFFF"
        for cell in row.cells:
            set_cell_shading(cell, bg)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = "Calibri"

    # Set table borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="{header_color}"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="{header_color}"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{header_color}"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="{header_color}"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{header_color}"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="{header_color}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def add_formatted_paragraph(doc, text, style="Normal", bold=False, alignment=None,
                            font_size=None, space_before=None, space_after=None,
                            font_name="Calibri", font_color=None):
    """Add a paragraph with formatting options."""
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.font.name = font_name
    if font_size:
        run.font.size = Pt(font_size)
    if bold:
        run.bold = True
    if font_color:
        run.font.color.rgb = font_color
    if alignment is not None:
        p.alignment = alignment
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p


def add_mixed_paragraph(doc, parts, style="Normal", alignment=None,
                        space_before=None, space_after=None):
    """Add a paragraph with mixed bold/normal text.
    parts is a list of tuples: (text, bold)
    """
    p = doc.add_paragraph(style=style)
    for text, bold in parts:
        run = p.add_run(text)
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.bold = bold
    if alignment is not None:
        p.alignment = alignment
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p


def add_bullet_list(doc, items, level=0):
    """Add bullet list items."""
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        # Handle mixed formatting in bullet items
        if isinstance(item, list):
            # item is a list of (text, bold) tuples
            for text, bold in item:
                run = p.add_run(text)
                run.font.name = "Calibri"
                run.font.size = Pt(11)
                run.bold = bold
        else:
            run = p.add_run(item)
            run.font.name = "Calibri"
            run.font.size = Pt(11)
        if level > 0:
            p.paragraph_format.left_indent = Inches(0.5 * (level + 1))


def add_sub_bullet_list(doc, items):
    """Add sub-bullet list items (indented)."""
    for item in items:
        p = doc.add_paragraph(style="List Bullet 2")
        if isinstance(item, list):
            for text, bold in item:
                run = p.add_run(text)
                run.font.name = "Calibri"
                run.font.size = Pt(11)
                run.bold = bold
        else:
            run = p.add_run(item)
            run.font.name = "Calibri"
            run.font.size = Pt(11)


def add_numbered_list(doc, items):
    """Add numbered list items."""
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph(style="List Number")
        if isinstance(item, list):
            for text, bold in item:
                run = p.add_run(text)
                run.font.name = "Calibri"
                run.font.size = Pt(11)
                run.bold = bold
        else:
            run = p.add_run(item)
            run.font.name = "Calibri"
            run.font.size = Pt(11)


def set_column_widths(table, widths):
    """Set column widths for a table."""
    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx < len(row.cells):
                row.cells[idx].width = Inches(width)


def add_page_break(doc):
    """Add a page break."""
    doc.add_page_break()


def setup_styles(doc):
    """Configure document styles."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(6)

    # Title style
    if "Title" in doc.styles:
        title_style = doc.styles["Title"]
        title_style.font.name = "Calibri"
        title_style.font.size = Pt(26)
        title_style.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
        title_style.font.bold = True

    # Heading 1
    if "Heading 1" in doc.styles:
        h1 = doc.styles["Heading 1"]
        h1.font.name = "Calibri"
        h1.font.size = Pt(18)
        h1.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
        h1.font.bold = True
        h1.paragraph_format.space_before = Pt(18)
        h1.paragraph_format.space_after = Pt(6)

    # Heading 2
    if "Heading 2" in doc.styles:
        h2 = doc.styles["Heading 2"]
        h2.font.name = "Calibri"
        h2.font.size = Pt(14)
        h2.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
        h2.font.bold = True
        h2.paragraph_format.space_before = Pt(12)
        h2.paragraph_format.space_after = Pt(4)

    # Heading 3
    if "Heading 3" in doc.styles:
        h3 = doc.styles["Heading 3"]
        h3.font.name = "Calibri"
        h3.font.size = Pt(12)
        h3.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
        h3.font.bold = True
        h3.paragraph_format.space_before = Pt(10)
        h3.paragraph_format.space_after = Pt(4)


def add_header_footer(doc):
    """Add header and footer to the document."""
    for section in doc.sections:
        # Set margins to 1 inch (2.54 cm)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

        # Header
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = hp.add_run("Fundaci\u00f3n Canguro - Propuesta Metodol\u00f3gica")
        run.font.name = "Calibri"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
        run.font.italic = True

        # Add a bottom border to header
        pPr = hp._p.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="2F5496"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)

        # Footer with page number
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add page number field
        run = fp.add_run()
        run.font.name = "Calibri"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._r.append(fldChar1)

        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run._r.append(instrText)

        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run._r.append(fldChar2)


def create_document():
    doc = Document()

    # Setup styles
    setup_styles(doc)

    # Setup header and footer
    add_header_footer(doc)

    # ==========================================
    # TITLE PAGE
    # ==========================================

    # Add some spacing before title
    for _ in range(4):
        doc.add_paragraph("")

    # Title
    title_p = doc.add_paragraph(style="Title")
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(
        "Propuesta Metodol\u00f3gica - Identificaci\u00f3n de Factores de Riesgo de "
        "Malnutrici\u00f3n en Ni\u00f1os Prematuros"
    )
    run.font.name = "Calibri"
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
    run.font.bold = True

    # Subtitle
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_before = Pt(12)
    run = subtitle_p.add_run(
        "Fundaci\u00f3n Canguro - Programa Madre Canguro Integral (PMCI)"
    )
    run.font.name = "Calibri"
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x59, 0x56, 0x59)
    run.font.italic = True

    # Add a decorative line
    line_p = doc.add_paragraph()
    line_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line_p.paragraph_format.space_before = Pt(24)
    run = line_p.add_run("_" * 60)
    run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

    # ==========================================
    # SECTION 1: Metodologia y Enfoque Propuesto
    # ==========================================
    add_page_break(doc)

    doc.add_heading("1. Metodolog\u00eda y Enfoque Propuesto", level=1)

    doc.add_heading("1.1 Enfoque T\u00e9cnico", level=2)

    doc.add_heading("Descripci\u00f3n de la propuesta de soluci\u00f3n", level=3)

    add_mixed_paragraph(doc, [
        ("El proyecto aborda dos problemas de aprendizaje autom\u00e1tico complementarios "
         "sobre datos tabulares longitudinales de ~64.801 historias cl\u00ednicas con 753 variables:", False)
    ])

    add_mixed_paragraph(doc, [
        ("Problema 1 - Clasificaci\u00f3n de estado nutricional a 12 meses de edad corregida: ", True),
        ("Construir modelos supervisados que clasifiquen a los sujetos en grupos de estado "
         "nutricional (Normal, Deficiente) seg\u00fan los tres indicadores definidos por la OMS:", False)
    ], space_before=6)

    add_bullet_list(doc, [
        [("Retraso en el crecimiento: ", True), ("z-score talla/edad (LAZ/HAZ) < \u22122", False)],
        [("Desnutrici\u00f3n aguda: ", True), ("z-score peso/talla (WLZ/WHZ) < \u22122", False)],
        [("Bajo peso: ", True), ("z-score peso/edad (WAZ) < \u22122", False)],
    ])

    add_mixed_paragraph(doc, [
        ("Problema 2 - Identificaci\u00f3n de factores de riesgo tempranos: ", True),
        ("Determinar qu\u00e9 variables de las fases iniciales del desarrollo (prenatal, nacimiento, "
         "hasta 40 semanas EG, 3 y 6 meses EC) son predictoras significativas de malnutrici\u00f3n "
         "a los 12 meses, con \u00e9nfasis en: ", False),
        ("RCIU, RCEU, velocidad de crecimiento, tipo de alimentaci\u00f3n y duraci\u00f3n de la "
         "posici\u00f3n canguro.", True)
    ], space_before=6)

    # --- Tecnicas de IA ---
    doc.add_heading("T\u00e9cnicas de IA y justificaci\u00f3n", level=3)

    table_data = [
        ["T\u00e9cnica", "Justificaci\u00f3n", "Alineaci\u00f3n con objetivos"],
        [
            "Gradient Boosting (XGBoost / LightGBM)",
            "Estado del arte en datos tabulares; manejo nativo de valores faltantes; alta capacidad predictiva",
            "Modelo principal de clasificaci\u00f3n para maximizar la detecci\u00f3n de malnutrici\u00f3n"
        ],
        [
            "Random Forest",
            "Buena interpretabilidad v\u00eda importancia de variables; robusto ante outliers y datos heterog\u00e9neos",
            "Modelo complementario y baseline robusto para comparaci\u00f3n"
        ],
        [
            "Regresi\u00f3n Log\u00edstica (L1/L2)",
            "Modelo base interpretable; los coeficientes indican direcci\u00f3n y magnitud del efecto de cada variable",
            "Baseline estad\u00edstico y referencia de interpretabilidad cl\u00ednica"
        ],
        [
            "SHAP (SHapley Additive exPlanations)",
            "Explicabilidad a nivel global e individual; identifica contribuci\u00f3n de cada variable a la predicci\u00f3n",
            "Directamente alineado con el objetivo de identificar factores de riesgo interpretables por neonat\u00f3logos"
        ],
        [
            "Clustering (K-Means, clustering jer\u00e1rquico)",
            "Agrupamiento no supervisado para descubrir perfiles de pacientes",
            "Apoya la hip\u00f3tesis de agrupar sujetos en niveles diferentes (Normal, Deficiente) de forma data-driven"
        ],
        [
            "Selecci\u00f3n de variables (Boruta, RFE, mutual information)",
            "Reducci\u00f3n de dimensionalidad de 753 variables a un subconjunto relevante",
            "Focaliza el an\u00e1lisis en las variables cl\u00ednicamente accionables"
        ],
    ]

    table = doc.add_table(rows=len(table_data), cols=3)
    table.style = "Table Grid"
    for i, row_data in enumerate(table_data):
        for j, cell_text in enumerate(row_data):
            table.rows[i].cells[j].text = cell_text
    set_column_widths(table, [2.0, 2.5, 2.5])
    format_table(table)

    # --- Analisis temporal ---
    doc.add_heading("An\u00e1lisis temporal adicional", level=3)

    add_mixed_paragraph(doc, [
        ("Se realizar\u00e1 el entrenamiento y evaluaci\u00f3n de modelos segmentados por ", False),
        ("periodos quinquenales (2005-2010, 2011-2015, 2016-2020, 2021-2025) ", True),
        ("para evaluar la evoluci\u00f3n de factores de riesgo y el impacto de cambios en "
         "protocolos de atenci\u00f3n del PMCI.", False)
    ])

    # --- Estrategia de modelado por fases ---
    doc.add_heading("Estrategia de modelado por fases", level=3)

    add_formatted_paragraph(doc,
        "Se entrenar\u00e1n modelos incrementales incorporando variables por fase temporal del desarrollo:")

    add_bullet_list(doc, [
        [("Modelo Fase 0-2: ", True), ("Solo variables prenatales y de nacimiento (predicci\u00f3n m\u00e1s temprana posible)", False)],
        [("Modelo Fase 0-3: ", True), ("+ variables hasta 40 semanas EG", False)],
        [("Modelo Fase 0-4: ", True), ("+ variables hasta 3 meses EC", False)],
        [("Modelo Fase 0-5: ", True), ("+ variables hasta 6 meses EC", False)],
        [("Modelo Fase 0-6: ", True), ("+ variables hasta 9 meses EC", False)],
    ])

    add_formatted_paragraph(doc,
        "Esto permite evaluar en qu\u00e9 momento del seguimiento se alcanza una capacidad "
        "predictiva cl\u00ednicamente \u00fatil.")

    # --- CRISP-DM ---
    doc.add_heading("Metodolog\u00eda de desarrollo: CRISP-DM", level=3)

    add_mixed_paragraph(doc, [
        ("Se adopta la metodolog\u00eda ", False),
        ("CRISP-DM ", True),
        ("(Cross-Industry Standard Process for Data Mining), est\u00e1ndar de la industria para "
         "proyectos de anal\u00edtica y machine learning:", False)
    ])

    add_numbered_list(doc, [
        [("Comprensi\u00f3n del negocio: ", True),
         ("Alineaci\u00f3n con los objetivos cl\u00ednicos de la Fundaci\u00f3n Canguro y "
          "definici\u00f3n formal de las variables objetivo.", False)],
        [("Comprensi\u00f3n de los datos: ", True),
         ("Exploraci\u00f3n del dataset de 64.801 registros y 753 variables, perfilamiento y "
          "an\u00e1lisis de calidad usando el diccionario de datos PhETI.", False)],
        [("Preparaci\u00f3n de datos: ", True),
         ("Limpieza, transformaci\u00f3n, ingenier\u00eda de caracter\u00edsticas, manejo de "
          "valores faltantes y codificaci\u00f3n de variables.", False)],
        [("Modelado: ", True),
         ("Entrenamiento, validaci\u00f3n cruzada y comparaci\u00f3n de modelos de clasificaci\u00f3n.", False)],
        [("Evaluaci\u00f3n: ", True),
         ("M\u00e9tricas de desempe\u00f1o (AUC-ROC, F1, Recall), an\u00e1lisis de interpretabilidad "
          "con SHAP y validaci\u00f3n con expertos neonat\u00f3logos.", False)],
        [("Despliegue: ", True),
         ("Registro de modelos en MLflow, dashboard interactivo para exploraci\u00f3n por expertos.", False)],
    ])

    # --- 1.2 Herramientas ---
    doc.add_heading("1.2 Herramientas, Bibliotecas y Plataformas", level=2)

    tools_data = [
        ["Categor\u00eda", "Herramienta", "Uso"],
        ["Lenguaje", "Python 3.11+", "Lenguaje principal de desarrollo"],
        ["Manipulaci\u00f3n de datos", "Pandas, NumPy",
         "Carga, limpieza, transformaci\u00f3n y an\u00e1lisis del dataset SPSS/XLSX"],
        ["Visualizaci\u00f3n", "Matplotlib, Seaborn, Plotly",
         "EDA, distribuciones, correlaciones, dashboards interactivos"],
        ["Machine Learning", "Scikit-learn",
         "Regresi\u00f3n log\u00edstica, Random Forest, preprocesamiento, validaci\u00f3n cruzada, m\u00e9tricas"],
        ["Gradient Boosting", "XGBoost, LightGBM",
         "Modelos principales de clasificaci\u00f3n sobre datos tabulares"],
        ["Interpretabilidad", "SHAP",
         "Explicabilidad de predicciones, identificaci\u00f3n de factores de riesgo"],
        ["Selecci\u00f3n de variables", "Boruta, Scikit-learn (RFE, mutual_info)",
         "Reducci\u00f3n de dimensionalidad de 753 a variables relevantes"],
        ["Manejo de desbalance", "imbalanced-learn (SMOTE, ADASYN)",
         "T\u00e9cnicas de sobremuestreo/submuestreo para clases minoritarias"],
        ["Tracking de experimentos", "MLflow",
         "Registro de m\u00e9tricas, par\u00e1metros e hiperpar\u00e1metros de cada experimento"],
        ["Notebooks", "Jupyter Notebook / JupyterLab",
         "Exploraci\u00f3n interactiva y documentaci\u00f3n reproducible"],
        ["Control de versiones", "Git + GitHub",
         "Versionamiento de c\u00f3digo y colaboraci\u00f3n del equipo"],
        ["Versionamiento de datos", "DVC (Data Version Control)",
         "Versionamiento del dataset"],
        ["Entorno", "Conda / pip + requirements.txt",
         "Gesti\u00f3n de dependencias reproducibles"],
    ]

    table = doc.add_table(rows=len(tools_data), cols=3)
    table.style = "Table Grid"
    for i, row_data in enumerate(tools_data):
        for j, cell_text in enumerate(row_data):
            table.rows[i].cells[j].text = cell_text
    set_column_widths(table, [1.8, 2.5, 2.7])
    format_table(table)

    # ==========================================
    # SECTION 2: Exploracion y Descripcion de los Datos
    # ==========================================
    add_page_break(doc)

    doc.add_heading("2. Exploraci\u00f3n y Descripci\u00f3n de los Datos", level=1)

    doc.add_heading("2.1 Caracter\u00edsticas de los Datos", level=2)

    doc.add_heading("Fuente de datos", level=3)

    add_mixed_paragraph(doc, [
        ("El dataset proviene de la base de datos cl\u00ednica de la ", False),
        ("Fundaci\u00f3n Canguro", True),
        (", que contiene historias cl\u00ednicas anonimizadas de pacientes del ", False),
        ("Programa Madre Canguro Integral (PMCI) ", True),
        ("recolectadas durante los \u00faltimos 25 a\u00f1os (aprox. 1993-2024).", False)
    ])

    # --- Perfilamiento ---
    doc.add_heading("Perfilamiento del dataset", level=3)

    profile_data = [
        ["Atributo", "Valor"],
        ["Registros (filas)", "64.801 historias cl\u00ednicas"],
        ["Variables (columnas)", "753"],
        ["Formato original", "SPSS exportado a XLSX"],
        ["Diccionario de datos",
         "PhETI con 634 variables documentadas, 7 hojas (README, Variables, Fases, "
         "Episodios, Temas de Inter\u00e9s, Prefijos)"],
        ["Estructura temporal",
         "Datos longitudinales: prenatal \u2192 nacimiento \u2192 40 sem EG \u2192 3m, 6m, 9m, 12m EC"],
    ]

    table = doc.add_table(rows=len(profile_data), cols=2)
    table.style = "Table Grid"
    for i, row_data in enumerate(profile_data):
        for j, cell_text in enumerate(row_data):
            table.rows[i].cells[j].text = cell_text
    set_column_widths(table, [2.0, 5.0])
    format_table(table)

    # --- Distribucion de variables ---
    doc.add_heading("Distribuci\u00f3n de variables por prefijo/categor\u00eda", level=3)

    prefix_data = [
        ["Prefijo", "Categor\u00eda", "Descripci\u00f3n"],
        ["IDEN", "Identificaci\u00f3n", "ID del sujeto, sede, fecha de parto"],
        ["CSP", "Condici\u00f3n social", "Situaci\u00f3n de pareja, vivienda, escolaridad, ingreso"],
        ["CP", "Control prenatal", "Peso/talla madre, IMC, alertas, medicamentos, edad materna"],
        ["PA", "Parto", "Lugar, tipo de parto, complicaciones, RPM"],
        ["ERN", "Reci\u00e9n nacido", "Peso, talla, per\u00edmetro cef\u00e1lico, sexo, Ballard, APGAR"],
        ["HD", "Hospitalizaci\u00f3n", "D\u00edas UCI, d\u00edas URN, ventilaci\u00f3n mec\u00e1nica, antibi\u00f3ticos"],
        ["AC", "Adaptaci\u00f3n canguro", "D\u00edas posici\u00f3n canguro, adaptaci\u00f3n al seno, lactancia"],
        ["CR", "Crecimiento", "z-scores de peso, talla, PC en cada punto temporal"],
        ["NUT", "Nutrici\u00f3n", "Tipo de alimentaci\u00f3n a 40sem, 3m, 6m, 9m, 12m"],
        ["NM", "Neuromotor", "INFANIB, evaluaciones de neurodesarrollo"],
        ["SEG", "Seguimiento", "Consultas, rehospitalizaciones, tamizaje"],
        ["OX", "Ox\u00edgeno", "Datos durante suministro de ox\u00edgeno"],
        ["PCANG", "Posici\u00f3n canguro", "Duraci\u00f3n y caracter\u00edsticas de la posici\u00f3n canguro"],
    ]

    table = doc.add_table(rows=len(prefix_data), cols=3)
    table.style = "Table Grid"
    for i, row_data in enumerate(prefix_data):
        for j, cell_text in enumerate(row_data):
            table.rows[i].cells[j].text = cell_text
    set_column_widths(table, [1.0, 2.0, 4.0])
    format_table(table)

    # --- Estadisticas clave ---
    doc.add_heading("Estad\u00edsticas clave del dataset", level=3)

    stats_data = [
        ["Variable", "Media", "Desv. Est.", "M\u00edn", "M\u00e1x"],
        ["ERN_Peso (peso al nacer, g)", "1.707,9", "412,4", "676", "2.900"],
        ["ERN_Ballard (edad gestacional, sem)", "32,9", "2,7", "25", "40"],
        ["HD_DiasUCI", "3,4", "8,2", "0", "74"],
        ["HD_TotalDiasHospital", "20,9", "17,9", "1", "107"],
        ["Ces\u00e1rea (%)", "77,3%", "\u2014", "0", "1"],
        ["PESO1500G (\u22641500g, %)", "31,8%", "\u2014", "0", "1"],
        ["Prematuro (%)", "92,0%", "\u2014", "0", "1"],
        ["MUERTE1ANO (%)", "2,3%", "\u2014", "0", "1"],
    ]

    table = doc.add_table(rows=len(stats_data), cols=5)
    table.style = "Table Grid"
    for i, row_data in enumerate(stats_data):
        for j, cell_text in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = cell_text
            # Center-align numeric columns
            if j > 0:
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_column_widths(table, [2.5, 1.0, 1.0, 0.8, 0.8])
    format_table(table)

    # ==========================================
    # 2.2 Problemas y Desafios Identificados
    # ==========================================
    doc.add_heading("2.2 Problemas y Desaf\u00edos Identificados", level=2)

    doc.add_heading("Calidad de datos", level=3)

    add_bullet_list(doc, [
        [("Codificaci\u00f3n de valores faltantes: ", True),
         ("Los valores faltantes est\u00e1n codificados como la cadena de texto ", False),
         ("#NULL! ", True),
         ("en lugar de valores nulos nativos. Esto impide que pandas los reconozca "
          "autom\u00e1ticamente y requiere conversi\u00f3n expl\u00edcita en el preprocesamiento.", False)],
        [("Alto porcentaje de datos faltantes: ", True),
         ("M\u00faltiples variables de condici\u00f3n social (CSP_TipoVivienda, CSP_IngresoMensual, "
          "CSP_NutricionFam), de control prenatal (CP_ARO, CP_MesInicCP) y de hospitalizaci\u00f3n "
          "presentan proporciones significativas de valores faltantes, probablemente por evoluci\u00f3n "
          "en los formularios de captura a lo largo de 25 a\u00f1os.", False)],
        [("Tipos de datos incorrectos: ", True),
         ("Muchas columnas num\u00e9ricas son le\u00eddas como tipo ", False),
         ("object ", True),
         ("debido a la presencia de las cadenas ", False),
         ("#NULL!", True),
         (", lo que requiere conversi\u00f3n de tipo despu\u00e9s de la limpieza.", False)],
        [("Valores inconsistentes: ", True),
         ("El diccionario indica ", False),
         ("-1 ", True),
         ("como valor de dato faltante (VAR-MISSING-VALUE), lo que podr\u00eda coexistir con ", False),
         ("#NULL! ", True),
         ("y otros indicadores.", False)],
    ])

    # --- Desbalance de clases ---
    doc.add_heading("Desbalance de clases", level=3)

    add_bullet_list(doc, [
        ("La variable objetivo (malnutrici\u00f3n a 12 meses: z-score < \u22122) representa una "
         "condici\u00f3n minoritaria en la poblaci\u00f3n. Se anticipa un desbalance significativo "
         "entre clases Normal vs. Deficiente, lo cual puede sesgar los modelos hacia la clase mayoritaria."),
        ("La mortalidad al primer a\u00f1o (2,3%) y la deserci\u00f3n a 12 meses generan p\u00e9rdida "
         "de sujetos que no llegan al punto de evaluaci\u00f3n final."),
    ])

    # --- Sesgos potenciales ---
    doc.add_heading("Sesgos potenciales", level=3)

    add_bullet_list(doc, [
        [("Sesgo de selecci\u00f3n: ", True),
         ("Solo incluye pacientes que ingresaron al PMCI; no es representativo de toda la "
          "poblaci\u00f3n de prematuros.", False)],
        [("Sesgo temporal: ", True),
         ("Los datos abarcan ~25 a\u00f1os, durante los cuales los protocolos de atenci\u00f3n "
          "y los est\u00e1ndares de registro han cambiado.", False)],
        [("Sesgo de supervivencia: ", True),
         ("Solo los sujetos que sobrevivieron hasta los 12 meses y que no desertaron del "
          "programa tienen datos completos de la variable objetivo.", False)],
        [("Sesgo geogr\u00e1fico: ", True),
         ("Los datos corresponden principalmente a sedes del PMCI en Colombia.", False)],
    ])

    # --- Informacion sensible ---
    doc.add_heading("Informaci\u00f3n sensible y restricciones", level=3)

    add_bullet_list(doc, [
        ("Los datos corresponden a historias cl\u00ednicas de pacientes pedi\u00e1tricos, "
         "informaci\u00f3n considerada altamente sensible bajo normativas de protecci\u00f3n de datos en salud."),
        ("La informaci\u00f3n est\u00e1 anonimizada (no contiene nombres, c\u00e9dulas ni datos "
         "de identificaci\u00f3n directa)."),
        ("Existe un acuerdo de confidencialidad firmado. El uso es exclusivo para este proyecto de investigaci\u00f3n."),
        ("Los datos residen en la Fundaci\u00f3n Canguro y su acceso est\u00e1 controlado por la misma organizaci\u00f3n."),
        ("Se requiere cumplir con principios de \u00e9tica en investigaci\u00f3n y protecci\u00f3n de datos de menores de edad."),
    ])

    # ==========================================
    # 2.3 Preparacion de Datos
    # ==========================================
    doc.add_heading("2.3 Preparaci\u00f3n de Datos", level=2)

    add_formatted_paragraph(doc,
        "Con base en la exploraci\u00f3n y los problemas identificados, se propone el siguiente "
        "pipeline de limpieza y preparaci\u00f3n:")

    # Paso 1
    doc.add_heading("Paso 1: Conversi\u00f3n de valores faltantes", level=3)

    add_bullet_list(doc, [
        [("Reemplazar todas las cadenas ", False),
         ("#NULL! ", True),
         ("y los valores ", False),
         ("-1 ", True),
         ("(seg\u00fan diccionario) por ", False),
         ("NaN ", True),
         ("nativo de pandas.", False)],
        ("Convertir columnas a sus tipos correctos (float, int, category) seg\u00fan el diccionario "
         "PhETI (VAR-TYPE-prim: BOOLEAN, INTEGER, FLOAT, ORDINAL, NOMINAL)."),
    ])

    # Paso 2
    doc.add_heading("Paso 2: An\u00e1lisis y tratamiento de datos faltantes", level=3)

    add_bullet_list(doc, [
        "Calcular el porcentaje de faltantes por variable y por fase temporal.",
        "Eliminar variables con >70% de datos faltantes que no sean cl\u00ednicamente cr\u00edticas.",
        "Imputaci\u00f3n por tipo de variable:",
    ])

    add_sub_bullet_list(doc, [
        [("Variables num\u00e9ricas continuas: ", True),
         ("imputaci\u00f3n por mediana (robusta ante outliers) o KNN Imputer para variables correlacionadas.", False)],
        [("Variables categ\u00f3ricas/ordinales: ", True),
         ("imputaci\u00f3n por moda o creaci\u00f3n de categor\u00eda \u201cDesconocido\u201d.", False)],
        [("Variables binarias (BOOLEAN): ", True),
         ("imputaci\u00f3n por moda dentro de subgrupos cl\u00ednicos.", False)],
    ])

    add_bullet_list(doc, [
        ("Crear indicadores binarios de \u201cdato faltante\u201d para variables donde la ausencia "
         "puede ser informativa."),
    ])

    # Paso 3
    doc.add_heading("Paso 3: Tratamiento de outliers", level=3)

    add_bullet_list(doc, [
        ("Identificar outliers en variables antropom\u00e9tricas (peso, talla, PC) usando rangos "
         "fisiol\u00f3gicamente plausibles definidos en el diccionario (VAR-MIN-VALUE, VAR-MAX-VALUE)."),
        ("Aplicar winsorizaci\u00f3n o eliminaci\u00f3n de valores biol\u00f3gicamente imposibles "
         "(ej: peso al nacer <400g o >5000g)."),
    ])

    # Paso 4
    doc.add_heading("Paso 4: Codificaci\u00f3n de variables", level=3)

    add_bullet_list(doc, [
        [("Variables nominales ", True),
         ("(ej: tipo de parto, sede): One-Hot Encoding o Target Encoding.", False)],
        [("Variables ordinales ", True),
         ("(ej: categor\u00edas de peso, edad gestacional): Ordinal Encoding respetando el orden.", False)],
        [("Variables booleanas: ", True),
         ("Codificaci\u00f3n binaria (0/1) ya presente en la mayor\u00eda.", False)],
    ])

    # Paso 5
    doc.add_heading("Paso 5: Ingenier\u00eda de caracter\u00edsticas", level=3)

    add_bullet_list(doc, [
        [("Construcci\u00f3n de la variable objetivo: ", True),
         ("Crear variables binarias para cada indicador de malnutrici\u00f3n a 12 meses "
          "(LAZ<\u22122, WLZ<\u22122, WAZ<\u22122) y una variable compuesta.", False)],
        [("Variables de velocidad de crecimiento: ", True),
         ("Aprovechar las variables de velocidad de z-score ya calculadas "
          "(velocidadzscorepeso40_3m, etc.) y crear deltas adicionales entre fases.", False)],
        [("Indicadores de RCIU/RCEU: ", True),
         ("Derivar o validar variables de retardo de crecimiento intrauterino y extrauterino.", False)],
        [("Duraci\u00f3n de lactancia materna exclusiva: ", True),
         ("Construir variable de duraci\u00f3n de LME a partir de ali40, ali3m, ali6m.", False)],
        [("Variables agregadas: ", True),
         ("N\u00famero total de complicaciones, score de riesgo social compuesto.", False)],
    ])

    # Paso 6
    doc.add_heading("Paso 6: Manejo de desbalance de clases", level=3)

    add_bullet_list(doc, [
        ("Evaluar la proporci\u00f3n de la clase minoritaria (malnutrici\u00f3n) despu\u00e9s de "
         "construir la variable objetivo."),
        ("Aplicar SMOTE (Synthetic Minority Over-sampling Technique) solo en el conjunto de entrenamiento."),
        ("Comparar con submuestreo aleatorio de la clase mayoritaria y pesos de clase (class_weight='balanced')."),
        ("Evaluar ADASYN como alternativa adaptativa a SMOTE."),
    ])

    # Paso 7
    doc.add_heading("Paso 7: Selecci\u00f3n de variables", level=3)

    add_bullet_list(doc, [
        ("Aplicar filtro inicial por correlaci\u00f3n con la variable objetivo (mutual information, chi-cuadrado)."),
        ("Usar Boruta (wrapper de Random Forest) para selecci\u00f3n robusta de variables relevantes."),
        ("Aplicar RFE (Recursive Feature Elimination) como t\u00e9cnica complementaria."),
        ("Agrupar variables por fase temporal para el an\u00e1lisis incremental por fases."),
    ])

    # Paso 8
    doc.add_heading("Paso 8: Partici\u00f3n del dataset", level=3)

    add_bullet_list(doc, [
        ("Divisi\u00f3n estratificada 70% entrenamiento / 15% validaci\u00f3n / 15% test, "
         "manteniendo la proporci\u00f3n de clases."),
        ("Validaci\u00f3n cruzada estratificada (Stratified K-Fold, k=5) para el proceso de "
         "selecci\u00f3n de modelos."),
        ("Partici\u00f3n adicional por periodos quinquenales para el an\u00e1lisis temporal."),
    ])

    # ==========================================
    # SECTION 3: Cronograma
    # ==========================================
    add_page_break(doc)

    doc.add_heading("3. Cronograma", level=1)

    add_formatted_paragraph(doc,
        "Cronograma tentativo para 8 semanas de trabajo con un equipo de 4 personas:")

    cronograma_data = [
        ["Semana", "Fase CRISP-DM", "Actividades", "Responsable(s)"],
        [
            "1",
            "Comprensi\u00f3n del negocio y datos",
            "Revisi\u00f3n detallada del diccionario PhETI y literatura m\u00e9dica. "
            "Reuni\u00f3n de alineaci\u00f3n con expertos de la Fundaci\u00f3n Canguro. "
            "Definici\u00f3n formal de variables objetivo y criterios de clasificaci\u00f3n.",
            "Todo el equipo"
        ],
        [
            "2",
            "Exploraci\u00f3n de datos (EDA)",
            "Carga y perfilamiento completo del dataset. An\u00e1lisis de distribuciones, "
            "valores faltantes, correlaciones. Documentaci\u00f3n de hallazgos de calidad. "
            "Generaci\u00f3n de visualizaciones exploratorias.",
            "2 personas: EDA.\n2 personas: configuraci\u00f3n del entorno (MLflow, DVC, repositorio)."
        ],
        [
            "3",
            "Preparaci\u00f3n de datos",
            "Limpieza: conversi\u00f3n de #NULL!, tipos de datos, outliers. "
            "Imputaci\u00f3n de valores faltantes. Codificaci\u00f3n de variables categ\u00f3ricas. "
            "Construcci\u00f3n de variable(s) objetivo.",
            "2 personas: pipeline de limpieza.\n2 personas: ingenier\u00eda de caracter\u00edsticas y variables derivadas."
        ],
        [
            "4",
            "Preparaci\u00f3n + Modelado inicial",
            "Finalizaci\u00f3n de feature engineering. Selecci\u00f3n de variables (Boruta, RFE). "
            "Manejo de desbalance de clases. Entrenamiento de modelos baseline "
            "(Regresi\u00f3n Log\u00edstica, Random Forest).",
            "2 personas: selecci\u00f3n de variables + desbalance.\n2 personas: modelos baseline + registro en MLflow."
        ],
        [
            "5",
            "Modelado avanzado",
            "Entrenamiento de XGBoost/LightGBM. Optimizaci\u00f3n de hiperpar\u00e1metros "
            "(Grid/Random Search o Bayesian). Modelos incrementales por fases temporales. "
            "Clustering no supervisado.",
            "2 personas: modelos de clasificaci\u00f3n.\n2 personas: clustering + an\u00e1lisis por fases."
        ],
        [
            "6",
            "Evaluaci\u00f3n",
            "Comparaci\u00f3n de modelos (AUC-ROC, F1, Recall). An\u00e1lisis SHAP global e "
            "individual. Identificaci\u00f3n de top factores de riesgo por fase. "
            "An\u00e1lisis temporal por quinquenios.",
            "2 personas: m\u00e9tricas y SHAP.\n2 personas: an\u00e1lisis temporal + visualizaciones de resultados."
        ],
        [
            "7",
            "Despliegue + Validaci\u00f3n",
            "Construcci\u00f3n de dashboard interactivo para exploraci\u00f3n por expertos. "
            "Preparaci\u00f3n de casos de uso para validaci\u00f3n con neonat\u00f3logos. "
            "Documentaci\u00f3n de la soluci\u00f3n.",
            "2 personas: dashboard y herramientas interactivas.\n2 personas: documentaci\u00f3n t\u00e9cnica y presentaci\u00f3n de resultados."
        ],
        [
            "8",
            "Entrega y presentaci\u00f3n",
            "Validaci\u00f3n final con expertos neonat\u00f3logos (1-2 casos de uso). "
            "Ajustes finales basados en retroalimentaci\u00f3n. Entrega de modelos registrados "
            "en MLflow, c\u00f3digo documentado y reporte final.",
            "Todo el equipo"
        ],
    ]

    table = doc.add_table(rows=len(cronograma_data), cols=4)
    table.style = "Table Grid"
    for i, row_data in enumerate(cronograma_data):
        for j, cell_text in enumerate(row_data):
            table.rows[i].cells[j].text = cell_text
    set_column_widths(table, [0.7, 1.5, 2.8, 2.0])
    format_table(table)

    # --- Hitos clave ---
    doc.add_heading("Hitos clave", level=3)

    hitos_data = [
        ["Hito", "Semana", "Entregable"],
        ["Kick-off y alineaci\u00f3n con expertos", "1", "Documento de requerimientos validado"],
        ["EDA completado", "2", "Reporte de exploraci\u00f3n y calidad de datos"],
        ["Dataset limpio y preparado", "3", "Pipeline de preparaci\u00f3n reproducible"],
        ["Modelos baseline entrenados", "4", "Resultados iniciales registrados en MLflow"],
        ["Mejor modelo seleccionado", "6", "Comparativa de modelos + factores de riesgo SHAP"],
        ["Dashboard funcional", "7", "Herramienta interactiva para expertos"],
        ["Entrega final", "8", "Reporte, modelos, c\u00f3digo y validaci\u00f3n con neonat\u00f3logos"],
    ]

    table = doc.add_table(rows=len(hitos_data), cols=3)
    table.style = "Table Grid"
    for i, row_data in enumerate(hitos_data):
        for j, cell_text in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = cell_text
            if j == 1:  # Center the "Semana" column
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_column_widths(table, [2.5, 1.0, 3.5])
    format_table(table)

    # ==========================================
    # SAVE
    # ==========================================
    output_path = (
        r"C:\Users\lfabi\OneDrive\Documents\developer\uniandes\desarrollo_proyectos"
        r"\proyecto_desarrollo_soluciones_canguro"
        r"\Propuesta_Metodologia_Fundacion_Canguro.docx"
    )
    doc.save(output_path)
    print(f"Document saved successfully to:\n{output_path}")
    print(f"File size: {os.path.getsize(output_path) / 1024:.1f} KB")


if __name__ == "__main__":
    create_document()
